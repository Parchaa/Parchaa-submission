"""
Generate messy, edge-case DOCX test documents for CDSCO RegAI — Set 2.
Deliberately uneven: some sections missing, PII buried in prose, ambiguous
classifications, duplicate pairs, mixed formatting.

Run: python3 generate_docs_2.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.dirname(os.path.abspath(__file__))


# ── Helpers ──────────────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    return doc

def h(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def p(doc, text, bold=False, italic=False, size=None, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def kv(doc, key, val):
    para = doc.add_paragraph()
    r1 = para.add_run(f"{key}: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = para.add_run(val)
    r2.font.size = Pt(10)
    return para

def rule(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

def tbl(doc, rows, cols, header=None, shade_hdr="D6E4F0"):
    t = doc.add_table(rows=1 if header else 0, cols=cols)
    t.style = 'Table Grid'
    if header:
        hrow = t.rows[0]
        for i, (cell, txt) in enumerate(zip(hrow.cells, header)):
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.bold = True
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), shade_hdr)
        for cell in hrow.cells:
            cell._tc.get_or_add_tcPr().append(shd)
    for row_data in rows:
        row = t.add_row()
        for cell, txt in zip(row.cells, row_data):
            cell.text = txt
    return t


# ─────────────────────────────────────────────────────────────────────────────
# DOC 1: Messy nephrology ward notes — tests anonymisation edge cases
# PII buried in prose, ALL-CAPS doctor names, multiple patients mentioned,
# partial identifiers, re-identification risks via clinical values
# ─────────────────────────────────────────────────────────────────────────────

def make_nephrology_ward_notes():
    doc = new_doc()

    h(doc, "NEPHROLOGY UNIT — WARD PROGRESS NOTES", level=1, color=(0, 70, 127))
    p(doc, "Govt. Medical College & Hospital, Nagpur | Dept. of Nephrology", italic=True)
    p(doc, "NOTE: These are working notes — not final discharge summary. Handwritten entries transcribed by intern.", bold=True, color=(160, 0, 0))
    rule(doc)

    # Deliberately messy — no table, PII all in running text
    p(doc, "PATIENT: Mohan Laxman Deshmukh, 52/M, admitted 04/02/2024 via casualty. R/O End Stage Renal Disease. Referred from Dr. PRAKASH GAWANDE (Akola, mob: 9422108734). Patient's wife Sushma Deshmukh (contact: 7758992341) accompanied. Aadhaar of patient: 6712 4521 9034. Address: House No. 14, Gandhi Nagar, Wardha Road, Nagpur — 440015. Employed at Nagpur Municipal Corporation (NMC Emp ID: NMC-2009-4421).", bold=False)

    doc.add_paragraph()
    p(doc, "ADMITTING DIAGNOSIS: CKD Stage 5 — likely secondary to uncontrolled T2DM (duration >15 yrs) + Hypertension. Patient reports seeing a 'local doctor' in Amravati (name unknown to patient) for past 5 years — no records available.")

    p(doc, "PMH: T2DM since approx 2008 — poor compliance. HTN since 2012. Cataract surgery — L eye, 2020 (Wockhardt Hospital Nagpur). No prior hospitalisation record available. Patient denies alcohol but family (son — Nikhil Deshmukh, 9765001234) states patient drinks regularly.")

    doc.add_paragraph()
    h(doc, "INVESTIGATION RESULTS (ADMISSION DAY)", level=2)
    # deliberately mixed style
    p(doc, "Serum Creatinine: 9.8 mg/dL (markedly elevated). Urea: 187 mg/dL. eGFR: 6 ml/min — consistent with ESRD. K+: 6.2 mEq/L — HYPERKALAEMIA, ECG changes noted (peaked T waves). Na+: 131. HbA1c: 11.4% — very poorly controlled diabetes. Hb: 7.1 g/dL — anaemia of CKD. Urine output last 24h: ~280 mL (oliguria).")
    p(doc, "USG Abdomen (04/02/2024): Both kidneys shrunken — R: 7.8 cm, L: 7.4 cm, echogenic — chronic changes. No obstruction.")
    p(doc, "2D Echo: EF 48%, mild LVH, Grade 1 diastolic dysfunction — likely uraemic cardiomyopathy. REPORT BY DR. VIJAY TALEKAR, MD (CARDIOLOGY).")

    doc.add_paragraph()
    h(doc, "MANAGEMENT — DAY BY DAY", level=2)
    p(doc, "04/02: Emergency haemodialysis initiated. Central line inserted R IJV by DR. SURESH NIMKAR, DM NEPHROLOGY (Reg no: MCI-2007-88341). Pre-dialysis K+ 6.2 → post 4.1 mEq/L. IV Furosemide 80mg BD. Restrict fluids 500ml/day + insensible losses. Phosphate binder (Sevelamer 800mg TDS) started. Erythropoietin 4000 IU SC 3x/week.")
    p(doc, "05/02: Patient more comfortable. Repeat K+ 4.3. Creatinine 8.9. Second dialysis session done. Dietitian counselling — wife Sushma attended (she can be reached at 7758992341 or sushmadeshmukh1971@yahoo.com). Started on Sodium Bicarbonate 650mg BD for acidosis correction.")
    p(doc, "06/02: Fever 38.4C — blood culture sent (report awaited). Central line site looks clean. Added Meropenem empirically (dose adjusted for renal function). Nephrology round by DR. SURESH NIMKAR — discussed AV fistula creation for long-term HD access. Patient and family counselled about ESRD prognosis — interpretor used (patient speaks Marathi only, wife speaks some Hindi).")
    p(doc, "07/02: Blood culture — Staph aureus (MSSA) bacteraemia. Sensitivity: Sensitive to Cloxacillin, Vancomycin. Changed to Cloxacillin 2g IV 6-hourly (dose adjusted). TTE — no vegetation seen. Duration of treatment: 14 days minimum. Social worker referral made — patient eligible for CM health scheme (Mahatma Jyotiba Phule Jan Arogya Yojana — patient card no: MJPJAY-MH-2023-7731209).")
    p(doc, "08/02: Third HD session. Urine output picking up slightly — 420 mL/24h. Fever subsiding. Cloxacillin continuing. AV fistula surgery scheduled for 15/02/2024 (Dr. Nimkar + vascular surgery team, Dr. ANIL BORKAR).")

    doc.add_paragraph()
    h(doc, "NOTE FROM INTERN (partially illegible transcription follows):", level=3)
    p(doc, "Pt. counselled re: long term dialysis — 3x/wk mandatory. Family v. stressed re: cost. Son Nikhil (9765001234) asking about govt. schemes. Referred to SW. FBS today: 198. Post lunch not done — patient refused food. Diet compliance poor per nursing notes.", italic=True)
    p(doc, "ALSO NOTE: The elderly male in Bed 4B (different patient — Ramkrishna Tupe, 74M, PIN: 440023, mob 8421076543) had his dialysis session confused with Deshmukh's slot on 06/02 — near-miss incident. Incident report filed (Ref: GMCH-NMI-2024-0089). Nursing supervisor Ms. KAVITA DESHPANDE informed.", italic=True)

    doc.add_paragraph()
    rule(doc)
    p(doc, "These notes are NOT to be shared outside clinical team. DPDP Act 2023 applies. Any queries — contact Dr. Suresh Nimkar, HOD Nephrology, Ext. 2241.", size=9, italic=True)

    path = os.path.join(OUT, "01_nephrology_ward_notes_ANONYMISE_THIS.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 2: Incomplete, ambiguous SAE report — tests classification edge cases
# Missing causality, vague narrative, multiple possible severity categories,
# some mandatory fields blank
# ─────────────────────────────────────────────────────────────────────────────

def make_incomplete_sae():
    doc = new_doc()

    h(doc, "INDIVIDUAL CASE SAFETY REPORT (ICSR)", level=1, color=(139, 0, 0))
    p(doc, "CIOMS Form — Partially Completed | DRAFT — Not Verified", italic=True, color=(160, 0, 0))
    rule(doc)

    kv(doc, "Case Reference", "SAE-2024-KA-0178 (DRAFT)")
    kv(doc, "Report Type", "Initial — 15-Day Expedited")
    kv(doc, "Date of Report", "12 March 2024")
    kv(doc, "Date of Event", "01 March 2024")
    kv(doc, "Date Sponsor Received", "05 March 2024")
    kv(doc, "Sponsor", "BioGenix Therapeutics India Pvt. Ltd., Bengaluru")
    kv(doc, "Drug", "Tocilizumab (RoActemra) 8mg/kg IV infusion — Study drug (biosimilar — BGX-TCZ-01)")
    kv(doc, "Protocol No.", "BGX-RA-2022-01 (Phase III RA study)")
    kv(doc, "Site", "St. John's Medical College, Bengaluru — Site IN-003")
    kv(doc, "PI", "Dr. Latha Krishnamurthy, MD, DM Rheumatology")

    doc.add_paragraph()
    h(doc, "Patient Information", level=2)
    kv(doc, "Patient ID", "BGX-RA-IN003-0019")
    kv(doc, "Age/Sex", "61F")
    kv(doc, "Weight", "54 kg")
    kv(doc, "Indication", "Moderate-to-Severe Rheumatoid Arthritis, inadequate response to MTX")
    kv(doc, "Medical History", "Type 2 DM, Hypertension, Hypothyroidism (on Levothyroxine 50mcg)")
    kv(doc, "Concomitant Meds", "Metformin 500mg BD, Amlodipine 5mg OD, Levothyroxine 50mcg OD, Folic Acid 5mg weekly")

    doc.add_paragraph()
    h(doc, "Event Narrative", level=2)
    p(doc, """Patient received 5th infusion of BGX-TCZ-01 (8mg/kg = 432mg IV) on 01 March 2024. Approximately 3 hours post-infusion, patient developed high grade fever (39.8°C), rigors, hypotension (BP 78/46 mmHg), tachycardia (HR 128/min). Infusion had been completed at that point.

Patient was given IV fluids (NS 500ml bolus), IV Paracetamol, and IV Hydrocortisone 200mg. Response was partial — BP improved to 92/60 within 30 minutes. Patient was admitted to ICU for monitoring.

ICU Course (01–04 March 2024):
— Blood cultures sent on admission. Results: No growth (5 days). Procalcitonin: 4.1 ng/mL (elevated). CRP: 187 mg/L.
— Chest X-ray: Bilateral lower zone haziness — ? early ARDS vs fluid overload.
— Ferritin: 8,400 ng/mL (severely elevated). IL-6 levels: Not done (lab not available at site).
— Impression: Infusion-related reaction vs Cytokine Release Syndrome (CRS) — cannot exclude early sepsis.
— Broad-spectrum antibiotics started (Piperacillin-Tazobactam + Amikacin). Continued for 5 days.
— ICU stay: 4 days. Transferred to general ward 05 March 2024. Discharged 08 March 2024.
— At discharge: BP 118/74, HR 88. Mild residual fatigue. No organ failure documented at discharge.

Outcome: Recovered with sequelae (fatigue persists — unable to perform daily activities without assistance as of report date 12 March 2024).""")

    doc.add_paragraph()
    h(doc, "Causality Assessment", level=2)
    p(doc, "PI Assessment: [NOT COMPLETED — PI unavailable at time of report preparation. To be updated in follow-up.]", color=(160, 0, 0))
    p(doc, "Sponsor Medical Monitor Assessment: Possible — temporal relationship with infusion. Cannot rule out CRS. Pending further information.", italic=True)
    kv(doc, "Expectedness", "Listed in IB v3.0 (mild-moderate infusion reactions) — Grade 3-4 CRS is NOT listed. Unexpected.")
    kv(doc, "Previous Similar Events in Study", "1 prior mild infusion reaction (Grade 1, self-limiting) at Site IN-001 in November 2023. This case is Grade 3-4.")

    doc.add_paragraph()
    h(doc, "Seriousness Criteria", level=2)
    p(doc, "[ ] Death\n[?] Life-Threatening — hypotension responded to treatment; unclear if patient was 'at immediate risk of death'\n[X] Hospitalisation Required — admitted to ICU 4 days\n[ ] Persistent Disability — fatigue present but expected to resolve\n[ ] Congenital Anomaly\n[?] Medically Important — CRS with ferritin 8400 may represent significant immune event\n\nNOTE: Primary reportable criterion selected as Hospitalisation Required. However, Life-Threatening status DEBATED — see internal memo dated 06/03/2024 (attached — NOT INCLUDED IN THIS SUBMISSION).")

    doc.add_paragraph()
    h(doc, "Sections Not Completed", level=2)
    p(doc, "The following mandatory fields are MISSING from this initial report:", color=(160, 0, 0))
    for item in [
        "PI causality assessment (PI travelling, return date unknown)",
        "Dechallenge/rechallenge assessment — drug permanently discontinued, rechallenge not planned",
        "Narrative of concomitant medication changes during event",
        "Laboratory data annexure (partial data only in narrative)",
        "DSMB notification status — unknown if notified",
    ]:
        doc.add_paragraph(item, style='List Bullet')
    p(doc, "This report is a draft. Full report will be submitted within 7 days. Please do not forward to CDSCO until verified.", bold=True, color=(160, 0, 0))

    rule(doc)
    p(doc, "Prepared by: Mr. Arun Kumar, PV Associate, BioGenix Therapeutics | Date: 12 March 2024", italic=True, size=9)

    path = os.path.join(OUT, "02_SAE_ambiguous_CRS_CLASSIFY_THIS.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 3 & 4: Semantic duplicate SAE pair — same event, different language
# Tests duplicate detection: cosine similarity + AI blended scoring
# Doc A: Clinical language, complete
# Doc B: Lay summary language, slightly different details
# ─────────────────────────────────────────────────────────────────────────────

def make_duplicate_sae_pair():
    # ── DOC A: Clinical version ──
    doc = new_doc()
    h(doc, "INDIVIDUAL CASE SAFETY REPORT — CASE A", level=1, color=(80, 0, 80))
    p(doc, "Reference: SAE-2024-TN-0291-A | Initial Report | Submitted: 18 February 2024", italic=True)
    rule(doc)

    kv(doc, "Drug", "Rituximab (MabThera) 375 mg/m² IV — Cycle 3")
    kv(doc, "Sponsor", "Roche Products India Pvt. Ltd.")
    kv(doc, "Study", "RIT-NHL-2022-02 — Phase II, Diffuse Large B-Cell Lymphoma")
    kv(doc, "Site", "Cancer Institute (WIA), Adyar, Chennai — Site IN-007")
    kv(doc, "Patient ID", "RIT-NHL-IN007-0033")
    kv(doc, "Age/Sex", "48M")
    kv(doc, "Date of Event", "09 February 2024")

    doc.add_paragraph()
    h(doc, "Clinical Narrative", level=2)
    p(doc, """Patient RIT-NHL-IN007-0033, a 48-year-old male with DLBCL (Stage III-B, Ann Arbor), received his third cycle of Rituximab infusion (375 mg/m² = 680mg) on 09 February 2024, commencing at 10:15 hrs. The infusion was initially proceeding at standard rate escalation.

At 11:40 hrs (approximately 85 minutes post-commencement, total dose administered ~450mg), the patient developed acute anaphylaxis characterised by:
— Urticaria: generalised, pruritic wheals across trunk and upper limbs
— Angioedema: lip swelling, periorbital oedema
— Bronchospasm: severe, SpO2 dropped to 83% on room air (baseline 98%)
— Hypotension: BP fell to 72/40 mmHg (baseline 124/78 mmHg)
— Tachycardia: HR 142 bpm
— Loss of consciousness: brief (approximately 30 seconds)

Infusion was IMMEDIATELY STOPPED. Emergency protocol activated. Adrenaline 0.5mg IM (thigh) administered at 11:42 hrs. IV antihistamine (Chlorpheniramine 10mg) and IV Hydrocortisone 200mg administered. High-flow O2 via face mask (15L/min). Patient transferred to resuscitation bay.

Response to adrenaline was rapid — BP improved to 94/60 within 8 minutes. SpO2 recovered to 96% on O2. Second dose of adrenaline not required. Patient monitored for 24 hours in HDU. No recurrence of anaphylaxis. Discharged next day in stable condition.

ICH E2A Seriousness: LIFE-THREATENING (patient was at immediate risk of death — loss of consciousness, SpO2 83%, BP 72/40).
Causality: CERTAIN — clear temporal relationship, known class effect, positive dechallenge, no alternative explanation.
Expectedness: LISTED in SmPC — anaphylaxis is a known ADR of Rituximab.""")

    kv(doc, "Outcome", "Recovered without sequelae")
    kv(doc, "Action Taken", "Rituximab permanently discontinued. Oncologist to determine alternative regimen.")
    kv(doc, "Regulatory Action", "7-Day expedited SUSAR report filed. DSMB notified 09 February 2024 (same day).")

    path = os.path.join(OUT, "03_SAE_anaphylaxis_CASE_A_DUPLICATE_CHECK.docx")
    doc.save(path)
    print(f"Saved: {path}")

    # ── DOC B: Lay language, slightly different framing ──
    doc = new_doc()
    h(doc, "ADVERSE EVENT REPORT — CASE B (FOLLOW-UP SUBMISSION)", level=1, color=(80, 0, 80))
    p(doc, "Reference: SAE-2024-TN-0291-B | Follow-Up Report | Submitted: 22 February 2024", italic=True)
    p(doc, "NOTE: This report was prepared by the site nurse coordinator from patient records. A separate medically reviewed ICSR (Case A) was submitted on 18 Feb 2024. This is a supplementary lay-language summary submitted at site's request.", italic=True)
    rule(doc)

    kv(doc, "Drug Name", "MabThera (Rituximab) — cancer treatment infusion, 3rd dose")
    kv(doc, "Hospital", "Cancer Institute WIA, Adyar, Chennai")
    kv(doc, "Patient Code", "IN007-0033 (48 year old male, DLBCL diagnosis)")
    kv(doc, "Incident Date", "9th February 2024, morning session")

    doc.add_paragraph()
    h(doc, "What Happened", level=2)
    p(doc, """During the patient's third chemotherapy infusion session, about an hour and twenty minutes after the drip was started, the patient had a severe allergic reaction. The patient's blood pressure dropped very low, he developed skin rashes all over his body, his lips became swollen, and he was having severe difficulty breathing. His oxygen levels dropped very low (shown as 83% on the pulse oximeter). The patient briefly lost consciousness for a short period.

The nursing staff immediately stopped the infusion and called the emergency team. The doctor on duty, Dr. Nirmala Selvam, gave an adrenaline injection immediately. The patient responded well to the emergency treatment — his blood pressure and oxygen levels came back to near-normal within about 10 minutes. He was shifted to the HDU ward for close monitoring overnight.

The patient was doing well the next day and was allowed to go home. He did not have any further reactions after the treatment.

The treating oncologist, Dr. Rajan Krishnan, has decided to stop this particular drug permanently and is looking at other treatment options for the patient's lymphoma.

The incident has been reported to the study sponsor (Roche) on the same day as per protocol requirements.""")

    kv(doc, "Severity", "Severe — life-threatening allergic reaction (anaphylaxis)")
    kv(doc, "Was the patient hospitalised?", "Yes — monitored overnight in HDU, discharged next day (10 February 2024)")
    kv(doc, "Recovered?", "Yes — full recovery, no lasting effects")
    kv(doc, "Was this expected?", "Allergic reactions to this drug are known, but this was very severe")
    kv(doc, "Drug stopped?", "Yes — permanently")

    rule(doc)
    p(doc, "Submitted by: Ms. Deepa Rao, Site Research Coordinator, Cancer Institute WIA | Date: 22 February 2024", italic=True, size=9)
    p(doc, "This is NOT a replacement for the medical ICSR filed separately. For official records, refer SAE-2024-TN-0291-A.", size=9)

    path = os.path.join(OUT, "04_SAE_anaphylaxis_CASE_B_DUPLICATE_CHECK.docx")
    doc.save(path)
    print(f"Saved: {path}")

    print("Saved: 03_SAE_anaphylaxis_CASE_A_DUPLICATE_CHECK.docx")
    print("Saved: 04_SAE_anaphylaxis_CASE_B_DUPLICATE_CHECK.docx")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 5: NDA application — several critical sections missing/incomplete
# Tests completeness module for New Drug Application scenario
# ─────────────────────────────────────────────────────────────────────────────

def make_nda_application():
    doc = new_doc()

    h(doc, "NEW DRUG APPLICATION — FORM 44", level=1, color=(0, 61, 122))
    p(doc, "As per Rule 21 read with Schedule Y, Drugs & Cosmetics Act 1940 (amended)", italic=True)
    p(doc, "Submitted to: CDSCO, New Delhi | Submission Date: 20 March 2024 | Mode: Physical + SUGAM Portal", bold=True)
    rule(doc)

    kv(doc, "Application No.", "NDA/2024/DL/0445 (Auto-generated)")
    kv(doc, "Application Type", "New Drug — Market Authorisation (Generic)")
    kv(doc, "Applicant", "Zydus Lifesciences Limited")
    kv(doc, "Registered Office", "Zydus Tower, Satellite Cross Roads, Ahmedabad — 380015, Gujarat")
    kv(doc, "Manufacturing Site", "Survey No. 396/4, Moraiya, Sarkhej-Bavla NH No. 8A, Ahmedabad — 382213")
    kv(doc, "Drug Licence", "GUJ/MFG/1995/0041 — Valid till 31 March 2026")
    kv(doc, "Contact", "Dr. Dharav Shah, Head Regulatory Affairs | 079-68665000 | regulatory@zyduslife.com")

    doc.add_paragraph()
    h(doc, "Part I: Drug / Product Details", level=2)
    drug_info = [
        ("Generic Name (INN)", "Metronidazole 0.75% w/w Topical Gel"),
        ("Brand Name (Proposed)", "MetroZyd Gel"),
        ("Dosage Form", "Topical Gel"),
        ("Strength", "0.75% w/w Metronidazole"),
        ("Pack Size", "30g tube, 15g tube"),
        ("Route of Administration", "Topical — Apply to affected area twice daily"),
        ("Therapeutic Category", "Dermatological — Antibacterial / Antiprotozoal"),
        ("ATC Code", "D06BX01"),
        ("Reference Listed Drug (RLD)", "MetroGel 0.75% (Galderma) — US FDA approved 1988, currently marketed"),
        ("Innovator Product Status", "Off-patent — Metronidazole patent expired globally"),
        ("Proposed Indication", "Topical treatment of inflammatory lesions (papules and pustules) of rosacea"),
        ("Target Population", "Adults ≥ 18 years"),
        ("International Approval Status", "Approved: USA (1988), EU (1995), UK, Canada, Australia"),
        ("Indian Approval Status", "Generic — not previously approved in India as 0.75% topical gel formulation"),
    ]
    tbl(doc, drug_info, 2, header=["Parameter", "Details"])

    doc.add_paragraph()
    h(doc, "Part II: Chemistry, Manufacturing & Controls (CMC)", level=2)
    h(doc, "2.1 Drug Substance (API)", level=3)
    p(doc, "Metronidazole API sourced from Aarti Drugs Ltd., Tarapur (DMF-IN-2017-0089). Active DMF on file with CDSCO. Certificate of Analysis (CoA) attached. WHO prequalified API supplier.")
    kv(doc, "Specification", "In-house specification ZL-API-MTZ-001 v2.0 aligned with IP 2022 monograph")
    kv(doc, "Analytical Method", "HPLC (validated) — Method Validation Report ZL-VAL-MTZ-002 attached")

    h(doc, "2.2 Drug Product (Formulation)", level=3)
    p(doc, "Formulation developed at Zydus R&D Centre, Ahmedabad. Pilot batches (3 × 10 kg) manufactured at manufacturing site. Batch manufacturing records for 3 pilot batches attached.")
    kv(doc, "Excipients", "Carbopol 980, Propylene Glycol, Edetate Disodium, Methylparaben, Purified Water")
    kv(doc, "Manufacturing Process", "Standard gel manufacturing — see SOP ZL-MFG-GEL-007 (attached)")

    h(doc, "2.3 Stability Data", level=3)
    p(doc, "Real-time stability: 6-month data available at 25°C/60%RH (ICH Zone II conditions). 12-month and 24-month data PENDING — study ongoing.", color=(160, 0, 0))
    p(doc, "Accelerated stability: 6 months at 40°C/75%RH — results attached. Assay at T6: 98.7% (spec: 97.0–103.0%). No significant degradation observed.", italic=False)
    p(doc, "CRITICAL GAP: Real-time stability data under Indian climatic conditions (ICH Zone IVb: 30°C/65%RH) is NOT available. Only ICH Zone II data submitted. India is a Zone IVb country — CDSCO typically requires Zone IVb stability data for topical formulations.", bold=True, color=(160, 0, 0))
    p(doc, "Proposed shelf life: 24 months. NOTE: Insufficient stability data to support 24-month claim — only 6-month real-time data available. Extrapolation applied per ICH Q1E but Zone IVb data absent.")

    doc.add_paragraph()
    h(doc, "Part III: Bioequivalence / Comparative Data", level=2)
    p(doc, "For topical products, FDA guidance (July 2018) permits in vitro skin permeation testing (IVPT) as alternative to clinical PK studies. IVPT study conducted at Zydus Research Centre.", italic=True)
    kv(doc, "Study Reference", "ZL-BE-MTZ-IVPT-001")
    kv(doc, "Test Product", "MetroZyd Gel 0.75% (Zydus)")
    kv(doc, "Reference Product", "MetroGel 0.75% (Galderma) — imported from US")
    kv(doc, "Method", "Franz diffusion cell — human cadaver skin — n=24 donors")
    kv(doc, "Results", "Flux ratio T/R: 0.94 (90% CI: 0.84–1.02) — within 80–125% equivalence limits")
    p(doc, "NOTE: IVPT acceptability for CDSCO is NOT yet confirmed. CDSCO has not published formal guidance on IVPT for topical generics as of application date. Clinical comparative trial data (small pilot study in rosacea patients) was requested at pre-submission meeting (November 2023) — this data is NOT SUBMITTED with this application.", bold=True, color=(160, 0, 0))

    doc.add_paragraph()
    h(doc, "Part IV: Clinical Data Summary", level=2)
    p(doc, "No new clinical trials conducted by Zydus. Literature-based justification submitted (bibliography of 14 published RCTs on metronidazole 0.75% gel in rosacea — attached as Appendix C).")
    p(doc, "Global clinical package: Reference to FDA NDA 19-617 (Galderma) — summary of pivotal trials attached.")
    p(doc, "Indian-specific data: NONE. No Phase III data in Indian population. Justification: Rosacea is not ethnicity-dependent based on literature review.", color=(160, 0, 0))

    doc.add_paragraph()
    h(doc, "Part V: Labelling", level=2)
    p(doc, "Draft Package Insert: ATTACHED — Version 0.1 (Draft). Note: Sections 4.4 (Special Warnings), 4.8 (Adverse Effects), and Section 5 (Pharmacological Properties) are INCOMPLETE — marked as 'TBD — to be finalised post CDSCO feedback'.", color=(160, 0, 0))
    p(doc, "Patient Information Leaflet: NOT SUBMITTED with this application. Will be provided after label finalisation.")
    p(doc, "Outer Carton Design: Draft attached. Hindi translation of key label elements: ABSENT.")

    doc.add_paragraph()
    h(doc, "Part VI: Document Checklist", level=2)
    checklist = [
        ("Form 44 (Application Form)", "SUBMITTED", "Signed by Head RA"),
        ("Covering letter", "SUBMITTED", ""),
        ("Drug Substance (API) DMF", "SUBMITTED", "DMF-IN-2017-0089, Aarti Drugs"),
        ("Formulation details & BMRs", "SUBMITTED", "3 pilot batches"),
        ("Analytical method validation", "SUBMITTED", "HPLC method"),
        ("Stability data — ICH Zone II, 6 months", "SUBMITTED", "Real-time and accelerated"),
        ("Stability data — ICH Zone IVb (India)", "NOT SUBMITTED", "CRITICAL GAP — data unavailable"),
        ("Stability data — 12/24 month real-time", "NOT SUBMITTED", "Study ongoing"),
        ("Bioequivalence — IVPT study", "SUBMITTED", "Acceptability with CDSCO uncertain"),
        ("Clinical pilot study (comparative)", "NOT SUBMITTED", "Requested at pre-sub meeting; not conducted"),
        ("Literature-based clinical justification", "SUBMITTED", "14 published studies"),
        ("Draft Package Insert", "PARTIALLY SUBMITTED", "Sections 4.4, 4.8, 5 incomplete"),
        ("Patient Information Leaflet", "NOT SUBMITTED", "To follow after label review"),
        ("Labelling — Hindi translation", "NOT SUBMITTED", ""),
        ("GMP Certificate — Manufacturing Site", "SUBMITTED", "Valid to March 2026"),
        ("CoA — API (Metronidazole)", "SUBMITTED", "WHO prequalified supplier"),
        ("Free Sale Certificate — USA", "SUBMITTED", "FDA approval NDA 19-617 (reference product)"),
        ("Risk Management Plan", "NOT SUBMITTED", "Applicant position: Not required for generic topical"),
        ("Manufacturing site inspection status", "SUBMITTED", "Last inspected November 2022 — CDSCO clearance letter attached"),
        ("Authorisation letter / Power of Attorney", "SUBMITTED", ""),
    ]
    tbl(doc, checklist, 3, header=["Document", "Status", "Remarks"])

    doc.add_paragraph()
    rule(doc)
    p(doc, "Declaration: All information provided is true and correct to the best of our knowledge.", bold=True)
    p(doc, "Signature: Dr. Dharav Shah, Head Regulatory Affairs, Zydus Lifesciences | Date: 20 March 2024", italic=True, size=9)

    path = os.path.join(OUT, "05_NDA_generic_metronidazole_COMPLETENESS_CHECK.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 6 & 7: Protocol comparison pair — oncology drug, meaningful amendments
# v1 vs v2 with significant changes that need flagging
# ─────────────────────────────────────────────────────────────────────────────

def make_protocol_comparison_pair():
    for is_v2 in [False, True]:
        doc = new_doc()
        ver = "2.0" if is_v2 else "1.0"
        date = "05 December 2023" if not is_v2 else "28 February 2024"
        fname = "07_oncology_protocol_v2_COMPARE_THESE.docx" if is_v2 else "06_oncology_protocol_v1_COMPARE_THESE.docx"

        h(doc, f"CLINICAL STUDY PROTOCOL — VERSION {ver}", level=1, color=(0, 61, 122))
        kv(doc, "Protocol No.", "ONC-PEM-2023-01")
        kv(doc, "Version", ver)
        kv(doc, "Date", date)
        kv(doc, "Drug", "Pembrolizumab (Keytruda) + Pemetrexed + Carboplatin")
        kv(doc, "Indication", "First-line treatment of metastatic non-squamous NSCLC")
        kv(doc, "Sponsor", "MSD Pharmaceuticals India Pvt. Ltd., Gurugram")
        kv(doc, "Phase", "Phase III (Bridging Study — Indian Population)")
        rule(doc)

        h(doc, "1. Objectives", level=2)
        h(doc, "1.1 Primary Objective", level=3)
        if not is_v2:
            p(doc, "To evaluate the overall survival (OS) benefit of pembrolizumab plus chemotherapy (pemetrexed + carboplatin) versus chemotherapy alone in treatment-naive patients with metastatic non-squamous NSCLC without EGFR/ALK alterations, in an Indian population cohort.")
        else:
            p(doc, "To evaluate the overall survival (OS) and progression-free survival (PFS) benefit of pembrolizumab plus chemotherapy (pemetrexed + carboplatin) versus chemotherapy alone in treatment-naive patients with metastatic non-squamous NSCLC without EGFR/ALK alterations, regardless of PD-L1 expression, in an Indian population cohort.")
            p(doc, "[AMENDMENT v2.0: PFS added as co-primary endpoint per CDSCO recommendation. PD-L1 expression removed as stratification factor — now all-comers regardless of PD-L1 TPS.]", italic=True, color=(139, 0, 0))

        h(doc, "1.2 Secondary Objectives", level=3)
        sec_v1 = [
            "Objective response rate (ORR) per RECIST 1.1",
            "Duration of response (DoR)",
            "Safety and tolerability — AE/SAE profile",
            "Patient-reported outcomes (PROs) using EORTC QLQ-C30",
        ]
        sec_v2 = [
            "Objective response rate (ORR) per RECIST 1.1",
            "Duration of response (DoR)",
            "Disease control rate (DCR) [NEW in v2.0]",
            "Safety and tolerability — AE/SAE profile, including immune-mediated AEs",
            "Patient-reported outcomes (PROs) using EORTC QLQ-C30 and LC13 module [LC13 added in v2.0]",
            "Biomarker analysis — tumour mutational burden (TMB) and PD-L1 expression as exploratory endpoints [NEW in v2.0 — moved from primary to exploratory]",
        ]
        for item in (sec_v2 if is_v2 else sec_v1):
            doc.add_paragraph(item, style='List Bullet')

        h(doc, "2. Study Design", level=2)
        kv(doc, "Design", "Randomised, Open-label, Active-controlled, Parallel-group, Multicentre")
        kv(doc, "Arms", "Arm A: Pembrolizumab 200mg IV Q3W + Pemetrexed 500mg/m² + Carboplatin AUC5 Q3W × 4 cycles, then Pembro + Pemetrexed maintenance")
        kv(doc, "", "Arm B: Pemetrexed 500mg/m² + Carboplatin AUC5 Q3W × 4 cycles, then Pemetrexed maintenance")
        if is_v2:
            kv(doc, "Treatment Duration", "Until progression, unacceptable toxicity, or 35 cycles (pembrolizumab) — maximum 2 years [AMENDED from 24 cycles in v1.0]")
        else:
            kv(doc, "Treatment Duration", "Until progression, unacceptable toxicity, or 24 cycles (pembrolizumab) — approximately 18 months")
        kv(doc, "Randomisation", "1:1, stratified by ECOG PS (0 vs 1) and smoking status (never vs former/current)")
        if is_v2:
            p(doc, "[AMENDMENT v2.0: Stratification factor changed — PD-L1 TPS stratum (high/low/negative) REMOVED. Now stratified by ECOG PS and smoking only. Rationale: Enrolment restricted to PD-L1 high in v1.0 was causing unacceptably slow recruitment.]", italic=True, color=(139, 0, 0))

        h(doc, "3. Study Population", level=2)
        h(doc, "3.1 Inclusion Criteria", level=3)
        inc_v1 = [
            "Age ≥ 18 years",
            "Histologically confirmed metastatic non-squamous NSCLC (Stage IV)",
            "No prior systemic therapy for metastatic disease",
            "EGFR wild-type and ALK-negative (local testing)",
            "PD-L1 TPS ≥ 1% (local or central testing)",
            "ECOG Performance Status 0 or 1",
            "Adequate organ function (haematologic, hepatic, renal, pulmonary)",
            "No active autoimmune disease requiring systemic treatment",
        ]
        inc_v2 = [
            "Age ≥ 18 years",
            "Histologically confirmed metastatic non-squamous NSCLC (Stage IV)",
            "No prior systemic therapy for metastatic disease",
            "EGFR wild-type and ALK-negative (local testing); ROS1 and MET exon 14 must also be negative [AMENDED — ROS1 and MET testing added v2.0]",
            "PD-L1 expression: Any level acceptable (PD-L1 testing no longer required for eligibility) [MAJOR AMENDMENT v2.0 — all-comers design]",
            "ECOG Performance Status 0 or 1",
            "Adequate organ function (haematologic, hepatic, renal, pulmonary)",
            "No active autoimmune disease requiring systemic treatment within past 2 years [AMENDED — added '2 years' qualifier]",
            "Creatinine clearance ≥ 45 mL/min (for pemetrexed dosing safety) [NEW criterion v2.0]",
        ]
        for item in (inc_v2 if is_v2 else inc_v1):
            doc.add_paragraph(item, style='List Bullet')

        h(doc, "3.2 Exclusion Criteria", level=3)
        excl_v1 = [
            "Squamous histology or mixed histology with predominant squamous component",
            "EGFR sensitising mutation (exon 19 del or L858R) or ALK rearrangement",
            "Prior treatment with anti-PD-1, anti-PD-L1, or anti-CTLA-4 antibody",
            "Active brain metastases (untreated or symptomatic)",
            "Interstitial lung disease or active pneumonitis",
            "Active TB or HBV/HCV infection",
        ]
        excl_v2 = [
            "Squamous histology or mixed histology with predominant squamous component",
            "EGFR sensitising mutation (exon 19 del or L858R), ALK rearrangement, ROS1 fusion, or MET exon 14 skipping [AMENDED]",
            "Prior treatment with anti-PD-1, anti-PD-L1, or anti-CTLA-4 antibody",
            "Untreated or symptomatic brain metastases — treated, stable brain mets now allowed [AMENDED — previously all brain mets excluded]",
            "Interstitial lung disease or active pneumonitis",
            "Active TB or HBV/HCV infection — latent TB allowed if prophylaxis initiated [AMENDED — latent TB exclusion relaxed given Indian TB burden]",
            "Chronic systemic corticosteroids > 10mg/day prednisone equivalent [NEW exclusion v2.0]",
        ]
        for item in (excl_v2 if is_v2 else excl_v1):
            doc.add_paragraph(item, style='List Bullet')

        h(doc, "4. Sample Size & Statistical Analysis", level=2)
        if not is_v2:
            p(doc, "Sample Size: 200 patients (100 per arm). Based on median OS 16 months (Pembro+Chemo) vs 11 months (Chemo), HR 0.60, 80% power, two-sided α 0.05, 10% dropout. Requires 120 events. Interim analysis: One at 60 events (50%) by IDMC.")
            p(doc, "Primary analysis: Stratified log-rank test. Kaplan-Meier estimates and Cox proportional hazards model. Interim analysis will use O'Brien-Fleming boundary.")
        else:
            p(doc, "[AMENDED v2.0]: Sample Size INCREASED to 300 patients (150 per arm). Rationale: (1) All-comers design (no PD-L1 filter) reduces expected treatment effect — revised HR 0.72. (2) Co-primary PFS endpoint requires additional events. Requires 180 OS events and 200 PFS events. 85% power. Dropout increased to 15% (recruitment from more diverse centres).")
            p(doc, "Interim analyses: TWO planned — at 40% and 70% events [increased from one interim in v1.0]. IDMC constituted — Dr. Ashok Datta (AIIMS, Statistics), Dr. Neha Krishnan (Oncology), independent statistician (Biostatistics Consulting Group, Pune).")
            p(doc, "Multiple testing strategy: Hierarchical — OS tested first; PFS tested only if OS significant at α 0.025. Prevents inflation of type I error.")

        h(doc, "5. Safety Monitoring & Reporting", level=2)
        p(doc, "SAEs must be reported to Sponsor within 24 hours of awareness. SUSAR reporting to CDSCO within 7 days (fatal/life-threatening) or 15 days (non-fatal). IDMC will review unblinded safety data at each interim analysis.")
        if is_v2:
            p(doc, "[AMENDMENT v2.0]: Enhanced immune-mediated AE monitoring protocol added — Grade 1 immune AEs must now be reported to site PI within 24 hrs; previously only Grade 2+ required reporting. New immune AE management algorithm (Appendix H) added. This was driven by hepatotoxicity cluster observation from another study site in Q4 2023.")

        h(doc, "6. Sites", level=2)
        sites_v1 = [
            ("IN-001", "Tata Memorial Hospital, Mumbai", "Dr. Kumar Prabhash"),
            ("IN-002", "AIIMS New Delhi", "Dr. Noopur Raje"),
            ("IN-003", "Kidwai Memorial Institute, Bengaluru", "Dr. Lokanatha D."),
            ("IN-004", "Cancer Institute WIA, Chennai", "Dr. Vijayalakshmi"),
        ]
        sites_v2 = sites_v1 + [
            ("IN-005", "HCG Oncology, Ahmedabad", "Dr. Chirag Desai [NEW SITE v2.0]"),
            ("IN-006", "Apollo Cancer Centre, Hyderabad", "Dr. Srikanth [NEW SITE v2.0]"),
        ]
        tbl(doc, sites_v2 if is_v2 else sites_v1, 3,
            header=["Site No.", "Institution", "Principal Investigator"])

        doc.add_paragraph()
        rule(doc)
        if is_v2:
            p(doc, "AMENDMENT SUMMARY (v1.0 → v2.0): Co-primary endpoint added (PFS); PD-L1 restriction removed; sample size increased 200→300; treatment duration extended 24→35 cycles; ROS1/MET testing added to eligibility; brain met exclusion relaxed; latent TB exclusion relaxed; 2 new sites added; enhanced immune AE monitoring protocol added.", bold=True, color=(139, 0, 0))
        p(doc, f"Approved by: Dr. Amit Chandra, Global Medical Officer, MSD India | Date: {date}", italic=True, size=9)

        path = os.path.join(OUT, fname)
        doc.save(path)
        print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 8: GCP site inspection notes — messy, unstructured
# Different from existing GMP inspection — tests report generation module
# Observations jumbled, severity not labelled, action items scattered in text
# ─────────────────────────────────────────────────────────────────────────────

def make_gcp_inspection():
    doc = new_doc()

    h(doc, "GCP SITE INSPECTION — WORKING NOTES", level=1, color=(80, 40, 0))
    p(doc, "FOR INTERNAL USE ONLY — DRAFT NOTES, NOT FOR DISTRIBUTION", bold=True, color=(160, 0, 0))
    rule(doc)

    kv(doc, "Inspection Type", "Announced GCP Site Inspection")
    kv(doc, "Study", "BGX-RA-2022-01 (Phase III RA Biosimilar Study — Tocilizumab)")
    kv(doc, "Site Inspected", "St. John's Medical College & Hospital, Bengaluru — Site IN-003")
    kv(doc, "Site Contact", "Dr. Latha Krishnamurthy, PI | Site Coordinator: Ms. Sunitha V. | 080-22065000")
    kv(doc, "Inspection Dates", "02–03 April 2024")
    kv(doc, "Lead Inspector", "Dr. Meena Iyer, Deputy Drugs Controller, CDSCO South Zone")
    kv(doc, "Co-Inspector", "Mr. Prakash Rao, Technical Officer, CDSCO")
    kv(doc, "Sponsor Representative", "Dr. Sunita Narayanan, CRA, BioGenix Therapeutics (observer)")
    doc.add_paragraph()

    h(doc, "Day 1 Observations (02 April 2024)", level=2)

    h(doc, "Informed Consent Process", level=3)
    p(doc, "OBS-ICF-01: Reviewed consent forms for all 22 enrolled subjects. Subject IN003-0019 (the subject who had the SAE in March 2024) — consent form reviewed. Date of consent: 14 Jan 2024. Consent was signed by subject. BUT — the witness signature field is blank. SOP requires independent witness for all oncology/serious disease studies. This was the case later involved in the CRS SAE.", color=(160, 0, 0))
    p(doc, "OBS-ICF-02: Subject IN003-0007 — consent form dated 22 November 2023. Corresponding study entry visit (Visit 1/Screening) date in CRF: 19 November 2023. Consent AFTER screening visit — informed consent obtained retrospectively, which is a GCP violation. Subject may have undergone screening procedures without consent.")
    p(doc, "OBS-ICF-03: ICF version used for subjects enrolled Nov 2023–Jan 2024 was v1.2 dated August 2023. HOWEVER, sponsor issued updated ICF v1.3 in October 2023 following a safety update (enhanced CRS warning). Site continued using v1.2 for 3 subjects enrolled after October 2023. These subjects never re-consented on updated form. Sponsor's communication to site (email dated 15 October 2023) requesting ICF update was acknowledged by coordinator but not actioned.", color=(160, 0, 0))

    h(doc, "Case Report Forms (CRFs) and Source Data", level=3)
    p(doc, "OBS-CRF-01: Electronic CRF system (Medidata Rave) audit trail reviewed. For subject IN003-0012, visit 4 (Week 12) data was entered on 25 February 2024. Visit date per CRF: 05 February 2024. 20-day delay in entry — exceeds 5-day entry window specified in DMP. CRF query raised by sponsor on 10 February 2024 (15 days earlier) — not responded to.")
    p(doc, "OBS-CRF-02: Discrepancy noted — source document (hospital case notes) for subject IN003-0017 records BP at Visit 3 as 148/92 mmHg. CRF entry for same visit shows 124/78 mmHg. When queried, coordinator stated 'I may have entered wrong subject's data'. No correction note made. This data was already used in a scheduled interim safety analysis submitted to sponsor in January 2024.")
    p(doc, "OBS-CRF-03: Laboratory results for infusion reaction workup of subject IN003-0019 (SAE case): original lab reports are printed email attachments — printouts faded, some values illegible. IL-6 levels — recorded in narrative as 'not done (lab not available)' but site lab has IL-6 capability per lab capability list. Clarification sought — coordinator unable to explain discrepancy during inspection.")

    h(doc, "Drug Accountability", level=3)
    p(doc, "OBS-DRUG-01: Investigational product (BGX-TCZ-01) stored in pharmacy cold room (2–8°C). Log reviewed — temperature within range throughout study period. Reconciliation done — 22 subjects × doses dispensed matches pharmacy record. OK.")
    p(doc, "OBS-DRUG-02: Destroyed drug documentation — 3 vials returned as partially used (following dose interruption for 3 subjects). Return-to-sponsor form completed but Destruction Certificate from central depot not available at site. Site expects to receive within 30 days.")
    p(doc, "OBS-DRUG-03: Emergency unblinding procedure — sealed code-break envelopes present. One envelope (Subject IN003-0019 — the SAE patient) has been opened and resealed with tape. Coordinator states code was broken during SAE management on 01 March 2024. However, unblinding documentation (Unblinding Record Form) was NOT completed. Only a handwritten note exists: 'code broken 1/3 emergency, Latha' — no timestamp, no formal record.")

    doc.add_paragraph()
    h(doc, "Day 2 Observations (03 April 2024)", level=2)

    h(doc, "Protocol Deviations", level=3)
    p(doc, "OBS-DEV-01: Protocol requires TB screening (Mantoux or IGRA) before first infusion. Reviewed all 22 subjects. Subject IN003-0009 — IGRA test done but result: 'Indeterminate'. Protocol and eligibility criteria state 'Negative IGRA required'. Subject was enrolled and received 6 infusions. This is a significant eligibility deviation — subject should not have been enrolled. Sponsor's medical monitor apparently knew (email from Dec 2023 in file) but no formal protocol deviation report submitted to CDSCO.")
    p(doc, "OBS-DEV-02: 4 subjects missed their Week 6 infusion (Visit 3) by more than the allowed window (±3 days). Protocol requires dose window management. These are listed as 'protocol deviations' in site deviation log. All 4 were reported to sponsor, but none were reported to the EC as required by St. John's EC SOP for deviations affecting subject safety.")
    p(doc, "OBS-DEV-03: Subject IN003-0022 — enrolled 20 March 2024. Lab results at screening show: ALT 78 U/L, AST 61 U/L. Protocol exclusion criterion: 'ALT or AST > 2.5× ULN' (ULN = 40 U/L). Subject's ALT is 1.95× ULN — within limit. However, Creatinine Clearance at screening: 43 mL/min (protocol minimum: 45 mL/min). This is a clear eligibility violation — subject did not meet minimum CrCl requirement. Enrolled despite this.")

    h(doc, "Ethics Committee (EC) Compliance", level=3)
    p(doc, "OBS-EC-01: Study approved by St. John's IEC — approval dated 15 September 2022. Annual renewal required. Renewal for 2024 (due September 2023): Certificate NOT available at site. Coordinator states application submitted but renewal not yet received — 7 months overdue. 9 subjects enrolled during period without valid EC renewal (October 2023 to present).")
    p(doc, "OBS-EC-02: SAE SAE-2024-KA-0178 (subject IN003-0019, CRS event March 2024) — EC notification of SAE was done on 15 March 2024 (14 days after event). St. John's EC SOP requires notification within 7 days. Delayed.")
    p(doc, "OBS-EC-03: Protocol amendment v2.0 (February 2024) — includes major changes (sample size, eligibility). Amendment submitted to EC: 01 March 2024. EC approval PENDING as of inspection date (02 April 2024). However, 2 new subjects were enrolled under v2.0 eligibility criteria on 15 March and 22 March 2024 BEFORE EC approval of the amendment. This is a GCP violation.")

    h(doc, "Site Staff & Training", level=3)
    p(doc, "OBS-TRN-01: GCP training records reviewed. PI (Dr. Krishnamurthy) — GCP certificate valid (2023). Co-Investigator Dr. Priya Menon — GCP certificate expired January 2024. No renewed certificate on file. Dr. Menon has been involved in Visit 3 and Visit 4 assessments for 6 subjects.")
    p(doc, "OBS-TRN-02: Protocol training sign-off for v2.0 — only 3 of 7 site staff have signed the training log for the amended protocol. Coordinator (Ms. Sunitha V.) — not signed. Sub-investigator Dr. Anjali Rao — not signed. Lab technician Ms. Preethi — not signed. Study procedures under v2.0 have been carried out by all staff regardless.")

    doc.add_paragraph()
    h(doc, "Inspector's Preliminary Classification of Findings", level=2)
    p(doc, "NOTE: This is a preliminary categorisation. Final classification subject to headquarters review.", italic=True)
    p(doc, "CRITICAL FINDINGS (require immediate CAPA + potential impact on data integrity / subject safety):", bold=True, color=(160, 0, 0))
    for obs in [
        "OBS-ICF-03: Subjects enrolled after safety update without re-consent on updated ICF",
        "OBS-ICF-02: Retrospective consent — screening performed before consent signed",
        "OBS-DEV-01: Ineligible subject (IGRA indeterminate) enrolled — sponsor aware but no formal deviation filed",
        "OBS-DEV-03: Eligibility violation — CrCl below minimum threshold",
        "OBS-EC-03: Subjects enrolled under unapproved protocol amendment",
        "OBS-EC-01: 9 subjects enrolled during period of lapsed EC approval",
    ]:
        doc.add_paragraph(obs, style='List Bullet')

    p(doc, "\nMAJOR FINDINGS:", bold=True, color=(200, 100, 0))
    for obs in [
        "OBS-CRF-02: Source data vs CRF discrepancy — may affect data integrity of interim safety analysis",
        "OBS-DRUG-03: Unblinding without formal documentation",
        "OBS-EC-02: Late SAE notification to EC",
        "OBS-TRN-01: Expired GCP certificate — sub-investigator conducting study procedures",
    ]:
        doc.add_paragraph(obs, style='List Bullet')

    p(doc, "\nMINOR FINDINGS:", bold=True, color=(100, 100, 0))
    for obs in [
        "OBS-ICF-01: Missing witness signature (subject IN003-0019)",
        "OBS-CRF-01: CRF data entry delays",
        "OBS-DRUG-02: Destruction certificate pending",
        "OBS-TRN-02: Protocol training sign-off incomplete for v2.0",
    ]:
        doc.add_paragraph(obs, style='List Bullet')

    doc.add_paragraph()
    rule(doc)
    p(doc, "Site response requested within 15 working days of formal report issuance. Data from this site will be flagged for sponsor review pending CAPA closure.", bold=True)
    p(doc, "Inspector: Dr. Meena Iyer | Date: 03 April 2024 | [Signature block — to be added to final report]", italic=True, size=9)
    p(doc, "Co-Inspector: Mr. Prakash Rao | Date: 03 April 2024", italic=True, size=9)

    path = os.path.join(OUT, "08_GCP_site_inspection_GENERATE_REPORT.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 9: Messy TEC meeting transcript — action items buried in prose
# Tests summarisation module (Meeting Transcript type)
# ─────────────────────────────────────────────────────────────────────────────

def make_messy_meeting_transcript():
    doc = new_doc()

    h(doc, "CDSCO — TECHNICAL EXPERT COMMITTEE (VACCINES & BIOLOGICALS)", level=1, color=(0, 61, 122))
    p(doc, "Subject Expert Committee Meeting — Biosimilars | 23rd Session | DRAFT MINUTES", bold=True)
    p(doc, "Date: 18 March 2024 | Venue: Conference Hall B, CDSCO HQ, FDA Bhavan, Kotla Road, New Delhi", italic=True)
    p(doc, "STATUS: DRAFT — Not ratified. Pending corrections from members. Do not cite.", bold=True, color=(160, 0, 0))
    rule(doc)

    p(doc, "The meeting was called to order at 10:20 AM (20 minutes behind schedule due to quorum issues — Dr. Sharma arrived late). Chair welcomed members and noted that Dr. Mathur (Pune) was joining via video call and Dr. Rajan could not attend as he is on leave following hospitalisation (get well soon wishes recorded).")

    h(doc, "Attendees", level=2)
    p(doc, "Present in person: Dr. R.K. Anand (Chair — Additional Drugs Controller), Dr. S. Krishnaswamy (Pharmacologist, JIPMER), Ms. Priya Mehta (CDSCO Deputy Controller — Biologicals), Dr. Y. Sharma (Immunologist, ICMR), Dr. Nalini Ghosh (Clinician — SSKM Kolkata), Mr. Tarun Bose (Legal Adviser, MoHFW)")
    p(doc, "Video call: Dr. Arvind Mathur (PGI Chandigarh, Endocrinology), Dr. (Ms.) Sunita Agarwal (Biostatistics, AIIMS Delhi)")
    p(doc, "Absent: Dr. P.K. Rajan (CMC Vellore) — on medical leave. Dr. H. Vohra (Industry representative, ABLE) — not invited to closed session items, joined only for Agenda Item 3.")
    p(doc, "Minutes: Mr. Kiran Shetty, Section Officer, CDSCO (no technical background — please verify all technical content before ratification)", italic=True)

    doc.add_paragraph()
    rule(doc)
    h(doc, "AGENDA ITEM 1: Application for Market Authorisation — Bevacizumab Biosimilar (Bevacirel) — Cipla Ltd.", level=2)
    p(doc, "[Presentation slides distributed physically — not attached to these minutes. Digital copy requested from Cipla but not yet received as of writing.]")

    doc.add_paragraph()
    p(doc, "10:25 AM — Dr. Anand introduced the application. Cipla's regulatory team (Dr. Rajesh Kumar, VP RA and Dr. Pooja Saxena, Medical Affairs) presented. This is Cipla's third biosimilar application this year.")

    p(doc, "Summary of Cipla Presentation: Bevacirel (bevacizumab biosimilar) — proposed for metastatic colorectal cancer, NSCLC, glioblastoma, and ovarian cancer (same indications as reference product Avastin, Roche). Analytical similarity exercise: 47 quality attributes tested. Functional assays: VEGF binding, ADCC, complement — results comparable within 90% CI for all critical quality attributes. PK comparability study: 54 healthy volunteers (single dose 1mg/kg IV), Indian subjects. PK parameters (AUC, Cmax, t½) bioequivalent to Avastin (all within 80–125%). Phase III efficacy study: 150 patients mCRC, randomised 1:1 Bevacirel vs Avastin (both with FOLFOX). Primary endpoint ORR: Bevacirel 38.7% vs Avastin 41.3% — equivalent (90% CI -12.1% to +7.1%, within equivalence margin ±15%). Immunogenicity: ADA positive at 24 weeks — Bevacirel 4.0%, Avastin 3.3% (no statistically significant difference). Safety profile similar.")

    doc.add_paragraph()
    p(doc, "10:55 AM — Committee discussion (verbatim notes — not cleaned up):", bold=True)

    p(doc, "Dr. Krishnaswamy raised the sample size issue immediately. 150 patients in Phase III is small for a biosimilar confirmation study. He noted that EMA guidelines for bevacizumab biosimilars recommended 300+ patients and that the CI for ORR was quite wide — the lower bound of -12.1% is clinically concerning if it represents real-world performance. He wanted a written response on why 150 was chosen and whether there was any power calculation justification. Cipla's Dr. Pooja Saxena said this was pre-agreed with CDSCO at a pre-submission meeting in April 2023 and referenced meeting minutes — Mr. Shetty to verify this with CDSCO files and report back to chair before next meeting.")

    p(doc, "Dr. Nalini Ghosh asked about GBM (glioblastoma) indication. Bevacizumab has an accelerated approval in US for GBM but this was subsequently withdrawn by FDA in 2023 after confirmatory trials failed to show OS benefit. She strongly objected to including GBM in the Indian label given this regulatory history. Dr. Anand agreed this needs to be flagged. Cipla team stated they had not been aware of the withdrawal. [Action: Cipla to submit updated global regulatory status document within 2 weeks — i.e. by 01 April 2024.]")

    p(doc, "Ms. Priya Mehta noted that biosimilar interchangeability — automatic substitution at pharmacy — is not currently part of the Indian regulatory framework but may become relevant as more biosimilars are approved. She suggested the committee issue a guidance statement. Dr. Anand agreed to bring this as a separate agenda item at the next meeting. Ms. Mehta to draft a 2-page note for pre-reading.")

    p(doc, "Dr. Sunita Agarwal (video call — some audio issues, she had to repeat parts of her comments) raised concerns about the immunogenicity data. The 24-week timepoint may be insufficient — bevacizumab antibodies can develop later. She recommended extending the immunogenicity follow-up to 52 weeks in the post-marketing study. The committee agreed this should be a condition of approval.")

    p(doc, "Dr. Sharma asked whether there is any Indian-specific hypersensitivity data. Infusion reactions are known with bevacizumab. Cipla confirmed 0 Grade 3-4 infusion reactions in the 150-patient study but Dr. Sharma noted the sample is too small to detect rare reactions reliably. He recommended a post-marketing pharmacovigilance registry of at least 500 patients with active follow-up for hypersensitivity.")

    doc.add_paragraph()
    p(doc, "12:00 PM — DECISION — Agenda Item 1:", bold=True, color=(0, 100, 0))
    p(doc, "After discussion, the committee voted 5:1 in favour of recommending conditional market authorisation for Bevacirel. Dr. Krishnaswamy was the dissenting vote — his primary concern was sample size adequacy. The conditions are as follows (Mr. Shetty please list carefully):")
    cond = [
        "GBM indication: EXCLUDED from Indian label pending review of FDA withdrawal and global data — Cipla to resubmit with updated data within 3 months",
        "Post-marketing Phase IV immunogenicity study: Extended to 52 weeks, minimum 200 patients, annual report to CDSCO",
        "Post-marketing pharmacovigilance registry: 500 patients, focus on hypersensitivity and infusion reactions, 18-month active follow-up",
        "Interchangeability: Label must explicitly state 'not automatically interchangeable' until CDSCO issues biosimilar interchangeability guidance",
        "Cipla to submit updated global regulatory status dossier (post-FDA GBM withdrawal) within 2 weeks",
    ]
    for c in cond:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph()
    rule(doc)
    h(doc, "AGENDA ITEM 2: Adalimumab Biosimilar Benefit-Risk Review — Post-Marketing Safety Signal", level=2)
    p(doc, "[Closed session — Industry representative Dr. Vohra was asked to wait outside. Joined after this item concluded at ~1:45 PM]")

    p(doc, "12:20 PM — Ms. Priya Mehta presented. Background: 3 approved adalimumab biosimilars in India (Exemptia, Adalimab-NK, Adipet). Post-marketing pharmacovigilance data pooled from all three — 12 cases of new-onset demyelinating disease (MS, optic neuritis) reported from 8 hospitals across India in past 18 months. Global rate for originator Humira: approx 0.1 per 1000 patient-years. Current Indian case count suggests higher-than-expected rate — however, causality is difficult to establish given comorbidities and lack of baseline neurological assessments in most patients.")

    p(doc, "Dr. Krishnaswamy: Given the known class effect of TNF-alpha inhibitors and demyelination (well-documented for infliximab and etanercept), this is a serious concern. India has high MS misdiagnosis rates — data quality is an issue. However, the precautionary principle applies.")

    p(doc, "Dr. Anand: CDSCO needs to take action. What is the minimum action that protects patients while not creating unnecessary panic?")

    p(doc, "After discussion — agreed actions (Dr. Anand to write formal order by 25 March 2024): All three manufacturers to issue Dear Healthcare Professional communication within 30 days — highlighting demyelination risk, contraindication in pre-existing MS/CNS demyelinating disease, requirement for baseline neurological assessment before initiation. Prescribing restriction — neurological evaluation required before starting adalimumab biosimilar in any patient with family history of MS. All three companies to submit updated PSUR with Indian demyelination data within 60 days.")

    doc.add_paragraph()
    rule(doc)
    h(doc, "AGENDA ITEM 3: Discussion — Regulatory Pathway for Complex Biosimilars (Monoclonal Antibody Mixtures)", level=2)
    p(doc, "[Dr. Vohra re-joined for this item at 1:48 PM. Also present: Dr. Vivek Rajan, Biocon (presenting)]")
    p(doc, "1:50 PM — Dr. Vivek Rajan (Biocon) presented industry perspective on regulatory clarity needed for bispecific antibody biosimilars — a category currently without any regulatory pathway in India. Countries like US and EU have not issued specific guidance either. Biocon flagged that 4 bispecific antibodies are coming off patent in 2026–2028 (catumaxomab, blinatumomab, emicizumab) and industry needs regulatory clarity now.")
    p(doc, "Dr. Anand: This is noted. CDSCO will form a working group. Timeline: working group to submit draft concept paper by December 2024. [Action: Ms. Priya Mehta to constitute working group, include industry and academic representatives, terms of reference by 15 April 2024]")
    p(doc, "Dr. Vohra (ABLE): Industry requests that any new guidelines go through a minimum 90-day public consultation. Committee agreed this is standard practice and would be followed.")

    doc.add_paragraph()
    rule(doc)
    h(doc, "MISCELLANEOUS", level=2)
    p(doc, "Dr. Anand reminded the committee that the CDSCO annual biosimilar report for FY 2023-24 is due by 30 April 2024. Ms. Mehta to compile. Dr. Ghosh mentioned that 3 queries from the previous meeting (22nd session — Denosumab biosimilar) are still pending responses from the sponsor (Natco Pharma). Mr. Shetty to follow up with Natco and report status to chair by 29 March 2024.")
    p(doc, "Next Meeting: Tentatively 22 April 2024, same venue. Confirmed once quorum confirmed — 5 of 8 members have travel conflicts that week. Ms. Priya Mehta to circulate Doodle poll by 25 March 2024.")
    p(doc, "Meeting closed approximately 3:10 PM.")
    rule(doc)
    p(doc, "Minutes recorded by: Mr. Kiran Shetty, Section Officer, CDSCO | Draft circulated: 20 March 2024", italic=True, size=9)
    p(doc, "IMPORTANT: These are unedited draft minutes. All technical details, figures, and decisions should be verified against presentation slides and audio recording before ratification.", bold=True, size=9)

    path = os.path.join(OUT, "09_TEC_biosimilar_meeting_SUMMARISE_THIS.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# Run all
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating test_docs_2 — messy, edge-case documents...\n")
    make_nephrology_ward_notes()
    make_incomplete_sae()
    make_duplicate_sae_pair()
    make_nda_application()
    make_protocol_comparison_pair()
    make_gcp_inspection()
    make_messy_meeting_transcript()
    print(f"\nAll 9 documents saved to: {OUT}")
    print("\nTesting guide:")
    print("  01_nephrology_ward_notes*.docx          → Anonymisation | PII buried in prose, ALL-CAPS doctors, re-ID risk via labs")
    print("  02_SAE_ambiguous_CRS*.docx              → Classification | Ambiguous: Life-Threatening vs Hospitalisation, missing causality")
    print("  03_SAE_anaphylaxis_CASE_A*.docx         → Duplicate Detection | Clinical language — anaphylaxis")
    print("  04_SAE_anaphylaxis_CASE_B*.docx         → Duplicate Detection | Lay language — same event, different words")
    print("  05_NDA_generic_metronidazole*.docx      → Completeness | NDA with critical gaps: no IVb stability, no clinical data")
    print("  06_oncology_protocol_v1*.docx           →")
    print("  07_oncology_protocol_v2*.docx           → Document Comparison | Many amendments: co-primary endpoint, sample size, eligibility")
    print("  08_GCP_site_inspection*.docx            → Inspection Report | GCP site (not GMP), ICF violations, protocol deviations")
    print("  09_TEC_biosimilar_meeting*.docx         → Summarisation (Meeting Transcript) | Action items buried in prose, messy format")
