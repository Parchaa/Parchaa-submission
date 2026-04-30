"""
Generate realistic DOCX test documents for each CDSCO RegAI module.
Run: python3 generate_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.dirname(os.path.abspath(__file__))


# ── Helpers ──────────────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    return doc


def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def para(doc, text, bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p


def kv(doc, key, value, indent=False):
    """Key: value paragraph."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run_k = p.add_run(f"{key}: ")
    run_k.bold = True
    run_k.font.size = Pt(10)
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)
    return p


def rule(doc):
    """Horizontal rule via paragraph border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row


def shade_row(row, hex_color="E8EDF5"):
    for cell in row.cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:color'), 'auto')
        shading_elm.set(qn('w:fill'), hex_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)


# ─────────────────────────────────────────────────────────────────────────────
# DOC 1: Patient Case Record (Anonymisation)
# ─────────────────────────────────────────────────────────────────────────────

def make_anonymisation_doc():
    doc = new_doc()

    # Header
    heading(doc, "APOLLO HOSPITALS — CLINICAL CASE SUMMARY", level=1, color=(0, 61, 122))
    para(doc, "Department of Cardiology & Internal Medicine | Apollo Hospitals, Delhi", italic=True)
    para(doc, "Confidential Patient Record — For Internal Use Only", bold=True, color=(180, 0, 0))
    rule(doc)

    # Patient demographics
    heading(doc, "Section 1: Patient Demographics", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    shade_row(t.rows[0], "D6E4F0")
    t.rows[0].cells[0].text = "Field"
    t.rows[0].cells[1].text = "Details"
    for row in t.rows:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True

    rows = [
        ("Patient Full Name",       "Ramesh Chandra Kumar"),
        ("Date of Birth",           "15/03/1978 (Age: 45 Years)"),
        ("Gender",                  "Male"),
        ("Aadhaar Number",          "2345 6789 0123"),
        ("PAN Card Number",         "ABCDE1234F"),
        ("Voter ID",                "XYZ1234567"),
        ("Passport Number",         "L2345678"),
        ("Mobile (Primary)",        "9876543210"),
        ("Mobile (Emergency)",      "8765432109"),
        ("Email Address",           "ramesh.kumar.1978@gmail.com"),
        ("Permanent Address",       "24, Rajpur Road, Dehradun, Uttarakhand — 248001"),
        ("Current Address",         "B-47, Sector 62, Noida, Uttar Pradesh — 201301"),
        ("Occupation",              "Senior Engineer, BHEL Ltd., Haridwar"),
        ("Religion",                "Hindu"),
        ("Blood Group",             "O+"),
        ("Marital Status",          "Married"),
        ("Spouse Name",             "Sunita Ramesh Kumar"),
        ("Spouse Contact",          "8765432109"),
        ("Emergency Contact",       "Vikram Kumar (Brother) — 9988776655"),
        ("Health Insurance",        "Star Health Comprehensive — Policy No. SHI-2024-00891234"),
        ("Employee ID (BHEL)",      "BHEL-HWR-04521"),
        ("Employer PAN",            "AAACB1234C"),
        ("Medical Record No.",      "IP/APL/DEL/2024/00456"),
        ("Referring Physician",     "Dr. Priya Sharma, MD (Internal Medicine), 9911223344"),
    ]
    for field, value in rows:
        r = table_row(t, [field, value], bold_first=True)

    doc.add_paragraph()

    # Admission details
    heading(doc, "Section 2: Admission & Treatment History", level=2)
    kv(doc, "Admission Date", "15 March 2024, 22:45 hrs (Emergency)")
    kv(doc, "Discharge Date", "23 March 2024, 11:00 hrs")
    kv(doc, "Ward", "Coronary Care Unit (CCU), Bed No. 4B")
    kv(doc, "Treating Physician", "Dr. Arvind Nair, DM (Cardiology), Reg. No. MCI-2003-45678")
    kv(doc, "Attending Nurse",    "Sr. Nurse Geeta Pillai, Staff No. APL-NSG-0892")
    kv(doc, "Anaesthesiologist",  "Dr. Siddharth Rao, MD (Anaesthesia)")

    doc.add_paragraph()
    heading(doc, "Section 3: Chief Complaints & History", level=2)

    para(doc, "Chief Complaints:", bold=True)
    for c in [
        "Acute onset severe chest pain (crushing, radiating to left arm and jaw) — duration 3 hours",
        "Profuse sweating and nausea for 2 hours prior to presentation",
        "Mild shortness of breath on exertion for past 2 weeks",
        "One episode of syncope (loss of consciousness) at home approximately 1 hour before arrival",
    ]:
        doc.add_paragraph(c, style='List Bullet')

    para(doc, "\nPast Medical History:", bold=True)
    for h in [
        "Type 2 Diabetes Mellitus — diagnosed February 2018, KEM Hospital Mumbai (Dr. S. Patil)",
        "Systemic Hypertension — on treatment since 2019",
        "HIV-positive (diagnosed March 2022, AIIMS Delhi) — on ART (Tenofovir + Lamivudine + Efavirenz)",
        "Chronic Kidney Disease Stage 2 (eGFR: 68 mL/min/1.73m²) — noted in 2023",
        "Coronary Artery Disease — diagnosed October 2023, PTCA to LAD performed at Fortis Gurgaon",
    ]:
        doc.add_paragraph(h, style='List Bullet')

    para(doc, "\nFamily History:", bold=True)
    doc.add_paragraph("Father: Deceased, myocardial infarction at age 58 — Mr. Shyam Lal Kumar, died 12 Jan 2002", style='List Bullet')
    doc.add_paragraph("Mother: Living, hypertension and T2DM — Mrs. Kamla Devi Kumar, DOB: 10/06/1950, Aadhar: 3456 7890 1234", style='List Bullet')
    doc.add_paragraph("Brother: Vikram Kumar (42M), healthy — contact 9988776655", style='List Bullet')

    doc.add_paragraph()
    heading(doc, "Section 4: Current Medications (Pre-Admission)", level=2)
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = 'Table Grid'
    shade_row(t2.rows[0], "D6E4F0")
    for cell, txt in zip(t2.rows[0].cells, ["Drug", "Dose", "Frequency", "Prescribing Doctor"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True

    meds = [
        ("Metformin 500mg",           "500mg",    "Twice daily",    "Dr. S. Patil, KEM Mumbai"),
        ("Telmisartan 40mg",          "40mg",     "Once daily",     "Dr. Priya Sharma, Apollo Delhi"),
        ("Atorvastatin 40mg",         "40mg",     "Bedtime",        "Dr. Priya Sharma, Apollo Delhi"),
        ("Aspirin 75mg",              "75mg",     "Once daily",     "Dr. Priya Sharma, Apollo Delhi"),
        ("Tenofovir/Lamivudine/EFV",  "Standard", "Once daily",     "Dr. Ankit Jha, AIIMS Delhi"),
        ("Clopidogrel 75mg",          "75mg",     "Once daily",     "Dr. Ravi Khanna, Fortis Gurgaon"),
        ("Pantoprazole 40mg",         "40mg",     "Before breakfast","Dr. Priya Sharma, Apollo Delhi"),
    ]
    for row in meds:
        table_row(t2, row)

    doc.add_paragraph()
    heading(doc, "Section 5: Investigations & Lab Reports", level=2)
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = 'Table Grid'
    shade_row(t3.rows[0], "D6E4F0")
    for cell, txt in zip(t3.rows[0].cells, ["Test", "Result", "Normal Range", "Date"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True

    labs = [
        ("Troponin I",            "18.6 ng/mL (HIGH)",  "< 0.04 ng/mL",       "15/03/2024"),
        ("CK-MB",                 "187 U/L (HIGH)",      "< 25 U/L",           "15/03/2024"),
        ("Blood Glucose (Fasting)","218 mg/dL (HIGH)",   "70–100 mg/dL",       "16/03/2024"),
        ("HbA1c",                 "9.2% (POORLY CTRL)", "< 7.0%",             "16/03/2024"),
        ("Serum Creatinine",      "1.6 mg/dL (HIGH)",   "0.7–1.3 mg/dL",      "16/03/2024"),
        ("eGFR",                  "58 mL/min (LOW)",     "> 90 mL/min",        "16/03/2024"),
        ("CD4 Count",             "412 cells/µL",        "> 500 cells/µL",     "16/03/2024"),
        ("HIV Viral Load",        "< 40 copies/mL",     "Undetectable target", "16/03/2024"),
        ("Total Cholesterol",     "228 mg/dL (HIGH)",   "< 200 mg/dL",        "17/03/2024"),
        ("LDL",                   "148 mg/dL (HIGH)",   "< 100 mg/dL",        "17/03/2024"),
        ("ECG",                   "ST elevation in V1–V4 (STEMI pattern)", "Normal sinus", "15/03/2024"),
        ("Echo (2D)",             "EF 38% (Reduced), Regional wall motion abnormality anterior wall", "EF > 55%", "16/03/2024"),
    ]
    for row in labs:
        table_row(t3, row)

    doc.add_paragraph()
    heading(doc, "Section 6: Diagnosis & Treatment Summary", level=2)
    para(doc, "Final Diagnosis:", bold=True)
    doc.add_paragraph("Acute ST-Elevation Myocardial Infarction (STEMI) — Anterior wall, LAD territory", style='List Bullet')
    doc.add_paragraph("Heart Failure with Reduced Ejection Fraction (HFrEF) — EF 38%", style='List Bullet')
    doc.add_paragraph("Type 2 Diabetes Mellitus — Poorly controlled", style='List Bullet')
    doc.add_paragraph("HIV infection — Virologically suppressed on ART", style='List Bullet')
    doc.add_paragraph("Chronic Kidney Disease Stage 3a", style='List Bullet')

    para(doc, "\nIntervention:", bold=True)
    para(doc, "Emergency Primary PCI performed on 16/03/2024 at 01:30 hrs. Drug-eluting stent (DES) placed in LAD (proximal). Procedure performed by Dr. Arvind Nair. Thrombus aspiration performed prior to stenting. Post-procedure TIMI 3 flow achieved. Patient shifted to CCU post-procedure for monitoring.")

    doc.add_paragraph()
    heading(doc, "Section 7: Discharge Instructions", level=2)
    kv(doc, "Follow-up", "Dr. Priya Sharma — 28 March 2024, 10:00 AM, Apollo Delhi OPD Room 12")
    kv(doc, "Cardiac Rehab", "Weekly sessions — Tuesdays and Fridays, Apollo Cardiac Rehab Centre")
    kv(doc, "Diabetologist", "Dr. S. Patil, KEM Mumbai — Teleconsult 29 March 2024")
    kv(doc, "Restriction", "No driving for 4 weeks. No strenuous activity for 8 weeks.")

    doc.add_paragraph()
    rule(doc)
    para(doc, "Prepared by: Dr. Arvind Nair | Apollo Hospitals Delhi | MCI Reg: 2003-45678 | Signed: 23/03/2024", italic=True, size=9)
    para(doc, "This document contains sensitive patient information. Unauthorized disclosure is an offence under DPDP Act 2023 and IPC Section 72.", bold=True, color=(180, 0, 0), size=9)

    path = os.path.join(OUT, "01_patient_case_record_ANONYMISE_THIS.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 2: SAE Case Narration
# ─────────────────────────────────────────────────────────────────────────────

def make_sae_narration_doc():
    doc = new_doc()

    heading(doc, "SERIOUS ADVERSE EVENT REPORT", level=1, color=(139, 0, 0))
    para(doc, "CIOMS Form I — Individual Case Safety Report (ICSR)", italic=True)
    para(doc, "Submitted to: Central Drugs Standard Control Organisation (CDSCO), New Delhi", bold=True)
    rule(doc)

    heading(doc, "A. Administrative Information", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    shade_row(t.rows[0], "F5E6E6")
    for cell, txt in zip(t.rows[0].cells, ["Field", "Value"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True

    admin = [
        ("Case Reference No.",       "SAE-2024-MH-0341"),
        ("Initial / Follow-up",      "Initial Report"),
        ("Date of Report",           "28 March 2024"),
        ("Date of Receipt by Sponsor","21 March 2024"),
        ("Date of Receipt by CDSCO", "28 March 2024 (within 7-day expedited reporting)"),
        ("Sponsor Name",             "Pharma Solutions India Pvt. Ltd., Mumbai"),
        ("Sponsor Address",          "Plot C-12, MIDC Andheri, Mumbai — 400093"),
        ("Sponsor Contact Person",   "Dr. Neha Srivastava, Head — Pharmacovigilance"),
        ("Sponsor Email",            "pvg@pharmasolutions.in"),
        ("Sponsor Phone",            "022-28456789"),
        ("Protocol Number",          "PSI-CV-2023-01, Version 4.0"),
        ("Study Title",              "A Phase III, Randomised, Double-Blind, Placebo-Controlled Study to Evaluate Efficacy and Safety of Cardivex 10mg in Adults with High Cardiovascular Risk"),
        ("IND/CTRI Number",          "CTRI/2023/06/054821"),
        ("EudraCT / NCT",            "NCT05612347"),
        ("Study Phase",              "Phase III"),
        ("Country of Occurrence",    "India"),
        ("Site Number",              "IN-008"),
        ("Site Name",                "KEM Hospital and Research Centre, Mumbai"),
        ("Principal Investigator",   "Dr. Anand Mehta, MD, DM (Cardiology)"),
        ("PI Contact",               "dr.anand.mehta@kemhospital.org | 022-24107000 Ext. 345"),
        ("Sub-Investigator",         "Dr. Pooja Kulkarni, MD"),
    ]
    for row in admin:
        table_row(t, row, bold_first=True)

    doc.add_paragraph()
    heading(doc, "B. Patient Information", level=2)
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Table Grid'
    shade_row(t2.rows[0], "F5E6E6")
    for cell, txt in zip(t2.rows[0].cells, ["Field", "Value"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True

    patient = [
        ("Patient ID (Blinded)",     "PSI-CV-IN008-0047"),
        ("Patient Full Name",        "Suresh Balakrishnan Iyer"),
        ("Date of Birth",            "22/07/1965"),
        ("Age at Enrolment",         "58 Years"),
        ("Gender",                   "Male"),
        ("Weight",                   "72 kg"),
        ("Height",                   "168 cm"),
        ("BMI",                      "25.5 kg/m²"),
        ("Aadhaar No.",              "4567 8901 2345"),
        ("Mobile",                   "9823401267"),
        ("Address",                  "Flat 3B, Swastik Park, Chembur, Mumbai — 400071"),
        ("Health Insurance",         "New India Assurance — Policy MH/2024/NIA/004512"),
        ("Date of Enrolment",        "14 November 2023"),
        ("Date of Randomisation",    "14 November 2023"),
        ("Treatment Arm",            "Active — Cardivex 10mg (double-blind)"),
        ("Screening No.",            "PSI-CV-SCR-IN008-0092"),
        ("Informed Consent Date",    "10 November 2023 (signed by patient and PI)"),
    ]
    for row in patient:
        table_row(t2, row, bold_first=True)

    doc.add_paragraph()
    heading(doc, "C. Relevant Medical History & Concomitant Medications", level=2)

    para(doc, "Pre-existing Conditions:", bold=True)
    conditions = [
        "Hypertension — diagnosed 2014, on treatment",
        "Hypercholesterolaemia — LDL 168 mg/dL at screening",
        "Anterior MI (non-STEMI) — June 2021, AIIMS Delhi, treated conservatively",
        "Type 2 Diabetes Mellitus — controlled, HbA1c 7.1% at screening",
        "Mild Chronic Obstructive Pulmonary Disease (COPD) — GOLD Stage 1",
        "Benign Prostatic Hyperplasia (BPH)",
    ]
    for c in conditions:
        doc.add_paragraph(c, style='List Bullet')

    para(doc, "\nConcomitant Medications at Time of Event:", bold=True)
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = 'Table Grid'
    shade_row(t3.rows[0], "F5E6E6")
    for cell, txt in zip(t3.rows[0].cells, ["Drug", "Dose", "Indication"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True
    cmeds = [
        ("Ramipril 5mg",          "5mg OD",        "Hypertension / Cardioprotection"),
        ("Aspirin 75mg",          "75mg OD",        "Antiplatelet"),
        ("Metformin 500mg",       "500mg BD",       "T2DM"),
        ("Tamsulosin 0.4mg",      "0.4mg OD",       "BPH"),
        ("Tiotropium inhaler",    "18µg OD",        "COPD"),
        ("Cardivex 10mg (Study)", "10mg OD",        "Study drug"),
    ]
    for row in cmeds:
        table_row(t3, row)

    doc.add_paragraph()
    heading(doc, "D. Event Description — Detailed Narrative", level=2)

    narrative = """On 21 March 2024 (Day 128 of treatment), the patient, Mr. Suresh Balakrishnan Iyer (58M, 72 kg), presented to the Emergency Department of KEM Hospital, Mumbai at 09:15 hrs with a 3-day history of progressively worsening bilateral lower limb weakness, difficulty rising from a sitting position, and inability to climb stairs. He reported dark brown (cola-coloured) urine since the morning of 21 March 2024. He denied any fever, joint pain, trauma, or recent vigorous exercise.

On examination, the patient was haemodynamically stable. There was marked proximal muscle weakness in bilateral lower limbs (MRC Grade 3/5). No skin rashes or arthritis noted. Urine dipstick showed 3+ blood with no red blood cells on microscopy — consistent with myoglobinuria.

EMERGENCY INVESTIGATIONS (21 March 2024):
- Serum Creatine Kinase (CK): 12,400 U/L (Normal: < 200 U/L) — 62× upper limit of normal
- Serum Creatinine: 2.1 mg/dL (Baseline at Screening: 0.98 mg/dL) — Acute Kidney Injury
- ALT: 312 U/L (HIGH), AST: 278 U/L (HIGH)
- Serum Myoglobin: 8,200 ng/mL (Normal: < 90 ng/mL)
- eGFR: 31 mL/min/1.73m² (Baseline: 84 mL/min/1.73m²)
- Blood Urea: 68 mg/dL (HIGH)
- Urine Myoglobin: Positive (strongly)
- Electrolytes: K+ 5.4 mEq/L (mildly elevated — risk of arrhythmia)
- ECG: Sinus rhythm, peaked T waves in precordial leads (hyperkalaemia pattern)

DIAGNOSIS: Statin-induced Rhabdomyolysis with Acute Kidney Injury (AKI) — classified as Hospitalisation Required (Serious) per ICH E2A criteria.

CLINICAL COURSE:
Day 1 (21 March): Study drug Cardivex 10mg immediately discontinued by PI. Patient admitted to Medical ICU. Aggressive IV hydration initiated (0.9% NaCl at 200 mL/hour). Urinary catheter placed — output monitored hourly. All nephrotoxic medications (Ramipril) withheld. Aspirin and Metformin continued.

Day 3 (23 March): CK trending down — 7,800 U/L. Urine output improving (>1 mL/kg/hour). Creatinine still elevated at 1.9 mg/dL. Patient transferred from ICU to general ward. Muscle weakness improving.

Day 5 (25 March): CK 2,100 U/L. Creatinine 1.4 mg/dL. Urine colour returning to normal. Patient ambulatory with assistance.

Day 8 (28 March): CK 380 U/L (approaching normal). Creatinine 1.1 mg/dL (near baseline). Patient discharged. Residual mild proximal muscle weakness noted at discharge — physiotherapy prescribed.

CAUSALITY ASSESSMENT:
The treating physician (Dr. Anand Mehta) and Sponsor Medical Monitor (Dr. Rohit Agarwal, MD) assessed causality as PROBABLE based on:
(1) Temporal relationship — event onset after 128 days of therapy with dose-consistent pattern
(2) Positive dechallenge — marked improvement after drug withdrawal
(3) No alternative explanation — no trauma, alcohol excess, or other causative drugs
(4) Known class effect — statin-induced rhabdomyolysis is a well-documented adverse drug reaction
(5) Lack of rechallenge (not attempted for patient safety reasons)

SERIOUSNESS CRITERIA (ICH E2A): Hospitalisation Required ✓ | Life-Threatening ✗ | Fatal ✗ | Persistent Disability ✗ | Congenital Anomaly ✗

REGULATORY HISTORY OF DRUG: Atorvastatin (active moiety of Cardivex) has 14 reported cases of rhabdomyolysis in FDA FAERS database. This is the first such case in the PSI-CV-2023-01 study. No previous safety signal identified in interim safety monitoring.

ACTIONS TAKEN:
1. Study drug permanently discontinued (21 March 2024)
2. DSMB notified on 21 March 2024 (same day)
3. All active study sites in India notified of potential signal via safety letter dated 23 March 2024
4. Protocol amendment proposed — CK monitoring to be added at monthly visits
5. IB to be updated with enhanced rhabdomyolysis warning

OUTCOME: Recovering — residual mild weakness at time of discharge. 30-day follow-up planned (20 April 2024).
FOLLOW-UP REPORT: Due within 15 days with 30-day CK and renal function results."""

    para(doc, narrative)

    doc.add_paragraph()
    heading(doc, "E. Regulatory Information", level=2)
    kv(doc, "Report Type",          "7-Day Expedited (Serious Unexpected Suspected ADR — SUSAR)")
    kv(doc, "Listedness",           "Unlisted — Rhabdomyolysis not in current IB (v5.0)")
    kv(doc, "Expectedness",         "Unexpected in context of this study")
    kv(doc, "Previous Similar SAEs","None in this study. 1 case of mild myalgia (non-serious) reported in March 2024.")
    kv(doc, "DSMB Review",          "Emergency session called 22 March 2024 — decision to continue study with enhanced monitoring")

    doc.add_paragraph()
    rule(doc)
    para(doc, "Prepared by: Dr. Neha Srivastava | Head Pharmacovigilance, Pharma Solutions India Pvt. Ltd.", italic=True, size=9)
    para(doc, "Verified by: Dr. Rohit Agarwal, MD | Medical Monitor | Date: 28 March 2024", italic=True, size=9)

    path = os.path.join(OUT, "02_SAE_narration_SUMMARISE_THIS.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 3: CDSCO Meeting Transcript
# ─────────────────────────────────────────────────────────────────────────────

def make_meeting_transcript_doc():
    doc = new_doc()

    heading(doc, "CENTRAL DRUGS STANDARD CONTROL ORGANISATION", level=1, color=(0, 61, 122))
    heading(doc, "TECHNICAL EXPERT COMMITTEE — ONCOLOGY & HAEMATOLOGY", level=2)
    para(doc, "MINUTES OF MEETING — 52nd Session", bold=True)
    rule(doc)

    kv(doc, "Date & Time", "15 April 2024, 10:00 AM — 4:30 PM")
    kv(doc, "Venue",       "Committee Room No. 2, CDSCO Headquarters, Kotla Road, New Delhi — 110002")
    kv(doc, "Chairperson", "Dr. V.G. Somani, Drugs Controller General of India (DCG(I))")
    kv(doc, "Minutes By",  "Mr. Ashish Verma, Section Officer, CDSCO")
    doc.add_paragraph()

    heading(doc, "Members Present", level=3)
    members = [
        ("Dr. V.G. Somani",       "DCG(I) — Chairperson"),
        ("Dr. Rajesh Bhatia",     "Professor of Medical Oncology, AIIMS New Delhi — Domain Expert"),
        ("Dr. Meena Rao",         "Head, Department of Pharmacology, CMC Vellore — Pharmacologist"),
        ("Dr. Sunil Kapoor",      "Senior Consultant, Haematology, PGI Chandigarh — Clinician"),
        ("Dr. Anita Singh",       "Joint Drugs Controller (JDC), CDSCO — Regulatory Expert"),
        ("Dr. Prashant Gupta",    "Additional Drugs Controller (ADC) — Oncology Division"),
        ("Dr. Kavitha Nair",      "Deputy Drugs Controller — Biologic Medicines"),
        ("Mr. Jayant Mishra",     "Legal Adviser, Ministry of Health & Family Welfare"),
        ("Dr. Sanjay Chandrasekhar","Industry Representative — OPPI (non-voting)"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    shade_row(t.rows[0], "D6E4F0")
    for cell, txt in zip(t.rows[0].cells, ["Name", "Designation", ""]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for name, desg in members:
        r = t.add_row()
        r.cells[0].text = name
        r.cells[1].text = desg

    heading(doc, "Members Absent (with apology)", level=3)
    doc.add_paragraph("Dr. H.G. Koshia (former DCG(I)) — Personal reasons", style='List Bullet')
    doc.add_paragraph("Dr. Shobha Bhatt (Paediatric Oncology, TATA Memorial) — Conference abroad", style='List Bullet')

    doc.add_paragraph()
    rule(doc)
    heading(doc, "AGENDA ITEM 1: Application for Expanded Indication of Pembrolizumab (Keytruda®) — MSD Pharmaceuticals", level=2)

    para(doc, "Application Reference: NDA/2023/MSD/0089 | Applicant: MSD Pharmaceuticals India Pvt. Ltd.\nIndication Sought: First-line treatment of metastatic non-small cell lung cancer (mNSCLC) without EGFR/ALK alterations, PD-L1 TPS ≥ 1%", italic=True)

    doc.add_paragraph()
    para(doc, "10:10 AM — Presentation by MSD Medical Director, Dr. Siddharth Mehrotra (MSD):", bold=True)
    para(doc, "Dr. Mehrotra presented data from KEYNOTE-024 and KEYNOTE-189 studies. Global data includes 2,843 patients across 189 sites in 32 countries. PFS improved from 6.0 months (chemotherapy) to 10.3 months (pembrolizumab, HR 0.50). OS benefit: 26.3 months vs 14.2 months. AE profile: 27.3% Grade 3-4 irAEs, manageable with standard protocols. FDA approved this indication in 2016; EMA in 2017. Pembrolizumab already approved in India for second-line NSCLC (2021).")

    para(doc, "\n10:45 AM — Questions from Committee:", bold=True)

    qa = [
        ("Dr. Rajesh Bhatia",
         "Indian patient data is insufficient. KEYNOTE-189 enrolled only 47 Indian patients out of 616 total (7.6%). Indian patients have different risk factors — higher proportion of non-smokers, different EGFR mutation profile. We cannot extrapolate Western data directly. Additionally, PD-L1 testing infrastructure at Indian tier-2/3 hospitals is inadequate — what is the plan?",
         "Dr. Mehrotra: Indian sub-group analysis shows HR of 0.52 (95% CI: 0.28–0.96), consistent with global data. Regarding PD-L1 testing — MSD commits to partnering with 200 government hospitals to establish testing capacity within 12 months of approval. We also propose a pharmacovigilance registry to collect Indian-specific real-world data."),
        ("Dr. Meena Rao",
         "The drug mechanism — PD-1 checkpoint inhibitor — carries risk of immune-mediated adverse events including pneumonitis, colitis, hepatitis. Indian population has higher TB burden. Is there data on safety in latent TB patients? ART drug interactions?",
         "Dr. Mehrotra: TB was an exclusion in KEYNOTE studies. However, real-world data from South Korea and Taiwan (similar TB burden) shows no significant increase in TB reactivation. We propose mandatory TB screening before initiation. No known PK interactions with standard anti-TB drugs — no CYP450 involvement for pembrolizumab."),
        ("Dr. Sunil Kapoor",
         "Cost is a major concern. At current MRP of ₹1,83,000 per vial, a full course costs ₹40–80 lakhs. This is inaccessible to 95% of Indian patients. Has the company considered differential pricing or patient assistance programmes?",
         "Dr. Sanjay Chandrasekhar (OPPI): Industry cannot comment on pricing in a regulatory meeting. However, MSD has an existing patient access programme — Keytruda PAP — providing free drug to BPL patients. Government may negotiate pricing under DPCO."),
        ("Dr. Anita Singh",
         "Under New Drugs Rules 2019, for drugs with global approval > 4 years and Indian data deficiency, we can mandate a post-approval commitment. I recommend: (1) Mandatory enrolment of 500 Indian patients in registry, (2) Annual PSUR submission, (3) Risk Management Plan with TB screening protocol.",
         "Committee agreed this is the appropriate regulatory pathway."),
    ]

    for questioner, question, answer in qa:
        para(doc, f"\n{questioner}:", bold=True)
        para(doc, f"Question: {question}", italic=True)
        para(doc, f"Response: {answer}")

    para(doc, "\n12:30 PM — Committee Deliberation (closed session):", bold=True)
    para(doc, "The committee deliberated for 45 minutes. Points raised: adequacy of Indian data, PD-L1 testing infrastructure gaps, cost accessibility, risk of immune-mediated adverse events in Indian population with higher infectious disease burden.")

    para(doc, "\nDECISION — Agenda Item 1:", bold=True, color=(0, 100, 0))
    decisions_1 = [
        "RECOMMENDATION: CONDITIONAL APPROVAL — Approved for first-line mNSCLC with PD-L1 TPS ≥ 50% (more restrictive than applied indication of ≥ 1%)",
        "Condition 1: Post-marketing pharmacovigilance registry — minimum 500 Indian patients, annual data submission to CDSCO for 5 years",
        "Condition 2: Indian-specific PK/PD study in 100 patients — to be submitted within 24 months",
        "Condition 3: Risk Minimisation Action Plan (RiMAP) including mandatory TB screening before initiation — to be submitted within 30 days",
        "Condition 4: Patient access programme mandatory — free drug for BPL patients (income < ₹3 lakh/year)",
        "Condition 5: Prescribing restricted to oncologists at NABH-accredited centres with immunotherapy management capability",
        "Label must include specific warning about TB reactivation risk",
    ]
    for d in decisions_1:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph()
    rule(doc)
    heading(doc, "AGENDA ITEM 2: Proposal — Relaxed Phase III Requirements for Rare Cancers (Orphan Drug Policy)", level=2)
    para(doc, "Presented by: Dr. Anita Singh (CDSCO) | Based on draft policy prepared by Oncology Division", italic=True)

    para(doc, "\n1:30 PM — Presentation:", bold=True)
    para(doc, "CDSCO proposes aligning with US Orphan Drug Act framework. Currently, all drugs require Phase III data regardless of disease prevalence. For conditions affecting < 1 in 10,000 population (rare cancers — e.g., cholangiocarcinoma, thymic carcinoma, SMARCB1-deficient tumours), conducting Phase III RCTs is impossible due to patient numbers. Proposed framework: Phase II single-arm data with validated biomarker endpoints (ORR, DoR) may be sufficient for Accelerated Approval with mandatory post-marketing confirmatory trials.")

    para(doc, "\n2:15 PM — Discussion:", bold=True)
    para(doc, "Dr. Bhatia raised concern that many rare cancer drugs entering with Phase II data have failed confirmatory trials (e.g., accelerated approvals withdrawn by FDA in 2022). The committee agreed that robust surrogate endpoints and mandatory confirmatory trial timelines are essential safeguards.")

    para(doc, "\nDECISION — Agenda Item 2:", bold=True, color=(0, 100, 0))
    for d in [
        "In-principle agreement to develop Rare Cancer Accelerated Approval pathway",
        "Sub-committee constituted: Dr. Bhatia (Chair), Dr. Singh, Dr. Kapoor",
        "Draft guidelines to be prepared by 30 June 2024",
        "Public consultation period: 60 days after draft release",
        "Target: Final guidelines published by 31 December 2024",
    ]:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph()
    rule(doc)
    heading(doc, "AGENDA ITEM 3: Safety Signal — Durvalumab + Tremelimumab Combination (HIMALAYA Study India Data)", level=2)
    para(doc, "Presented by: Dr. Kavitha Nair", italic=True)

    para(doc, "\nDr. Nair presented 6 cases of Grade 4 immune-mediated hepatitis from 3 Indian sites. Global incidence is 3.7% but Indian sites showing 8.9% (suspected underreporting in other sites or genuine ethnic difference). AZ committed to root cause analysis within 45 days.")

    para(doc, "\nDECISION — Agenda Item 3:", bold=True, color=(0, 100, 0))
    for d in [
        "Urgent Dear Healthcare Professional letter to be issued within 7 days",
        "Enhanced liver function monitoring: LFTs at baseline, Week 2, Week 4, then monthly",
        "AstraZeneca to submit root cause analysis within 45 days",
        "DSMB review of all Indian patients on combination by 30 April 2024",
    ]:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph()
    rule(doc)
    heading(doc, "ACTION ITEMS SUMMARY", level=2)
    t_ai = doc.add_table(rows=1, cols=4)
    t_ai.style = 'Table Grid'
    shade_row(t_ai.rows[0], "D6E4F0")
    for cell, txt in zip(t_ai.rows[0].cells, ["#", "Action", "Owner", "Deadline"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True
    ais = [
        ("1", "Issue RiMAP for Pembrolizumab including TB screening protocol", "MSD Pharmaceuticals", "15 May 2024"),
        ("2", "Issue conditional approval letter for Pembrolizumab", "Dr. Anita Singh / CDSCO", "30 April 2024"),
        ("3", "Draft guidelines for Orphan Drug Accelerated Approval", "Dr. Bhatia Sub-committee", "30 June 2024"),
        ("4", "Issue Dear HCP letter — Durvalumab hepatitis safety signal", "Dr. Kavitha Nair / CDSCO", "22 April 2024"),
        ("5", "Root cause analysis for Indian hepatitis cluster", "AstraZeneca India", "30 May 2024"),
        ("6", "Circulate approved minutes for ratification", "Mr. Ashish Verma", "22 April 2024"),
    ]
    for row in ais:
        table_row(t_ai, row)

    doc.add_paragraph()
    kv(doc, "Next Meeting", "20 May 2024, 10:00 AM, Committee Room 2, CDSCO HQ")
    kv(doc, "Meeting Closed", "4:30 PM")

    rule(doc)
    para(doc, "Minutes prepared by: Mr. Ashish Verma, Section Officer | Reviewed by: Dr. Anita Singh, JDC", italic=True, size=9)
    para(doc, "DRAFT — Subject to ratification at next meeting", bold=True, size=9)

    path = os.path.join(OUT, "03_CDSCO_meeting_transcript_SUMMARISE_THIS.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 4: SUGAM Clinical Trial Application (Completeness Check)
# ─────────────────────────────────────────────────────────────────────────────

def make_sugam_application_doc():
    doc = new_doc()

    heading(doc, "FORM CT-04 — APPLICATION FOR PERMISSION TO CONDUCT CLINICAL TRIAL", level=1, color=(0, 61, 122))
    para(doc, "As per Rule 101 of New Drugs and Clinical Trials Rules, 2019", italic=True)
    para(doc, "Submitted via CDSCO SUGAM Portal | Application Date: 01 April 2024", bold=True)
    rule(doc)

    kv(doc, "Application Reference No.", "CT-2024-MH-8821 (System Generated)")
    kv(doc, "Submission Mode",           "Online — SUGAM Portal (sugam.gov.in)")
    kv(doc, "Application Type",         "New Clinical Trial Application — Phase III")

    heading(doc, "Part A: Applicant Information", level=2)
    applicant = [
        ("Applicant Name",         "Sun Pharmaceutical Industries Limited"),
        ("Registered Address",     "SPARC, Tandalja, Vadodara, Gujarat — 390020"),
        ("Manufacturing Address",  "Plot 107-A, MIDC Silvassa, Dadra & Nagar Haveli — 396230"),
        ("CIN",                    "L24230GJ1993PLC019050"),
        ("GSTIN",                  "24AAACS9690N1ZP"),
        ("Drug Licence No.",       "MFG/GUJ/2001/0892 (Valid up to 31 Dec 2025)"),
        ("Contact Person",         "Dr. Anil Varghese, VP Clinical Development"),
        ("Email",                  "clinicaltrials@sunpharma.com"),
        ("Phone",                  "022-43240000"),
        ("Regulatory Affairs Head","Ms. Ritu Patel, Director Regulatory Affairs"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    shade_row(t.rows[0], "D6E4F0")
    for cell, txt in zip(t.rows[0].cells, ["Field", "Details"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in applicant:
        table_row(t, row, bold_first=True)

    doc.add_paragraph()
    heading(doc, "Part B: Drug / Investigational Product Details", level=2)
    drug = [
        ("Drug Name (INN)",         "Empagliflozin + Metformin Hydrochloride (Fixed Dose Combination)"),
        ("Brand Name (Proposed)",   "EmpaMet Plus"),
        ("Dosage Form",             "Film-coated Tablet"),
        ("Strengths",               "Empagliflozin 10mg + Metformin 500mg (Arm 1); Empagliflozin 25mg + Metformin 1000mg (Arm 2)"),
        ("Route of Administration", "Oral"),
        ("Therapeutic Class",       "Antidiabetic — SGLT2 Inhibitor + Biguanide FDC"),
        ("Mechanism of Action",     "Empagliflozin: Selective SGLT2 inhibitor — reduces renal glucose reabsorption. Metformin: Biguanide — reduces hepatic glucose output, improves insulin sensitivity"),
        ("ATC Code",                "A10BD20"),
        ("Drug Master File (DMF)",  "API-1: Empagliflozin DMF-US-034521 (Boehringer Ingelheim); API-2: Metformin DMF-IN-2019-0091"),
        ("Shelf Life",              "24 months from date of manufacture"),
        ("Storage Conditions",      "Store below 30°C, protect from moisture"),
        ("Sponsor Country of Origin","Germany / India (technology transfer)"),
        ("Global IND Status",       "Active IND 125430 (US FDA); IMPD submitted to EMA"),
    ]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Table Grid'
    shade_row(t2.rows[0], "D6E4F0")
    for cell, txt in zip(t2.rows[0].cells, ["Parameter", "Details"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in drug:
        table_row(t2, row, bold_first=True)

    doc.add_paragraph()
    heading(doc, "Part C: Study Summary", level=2)
    kv(doc, "Study Title",     "A Phase III, Randomised, Double-Blind, Active-Controlled Study to Evaluate the Efficacy and Safety of Empagliflozin/Metformin FDC Compared to Metformin Monotherapy as Add-on to Diet and Exercise in Adults with Type 2 Diabetes Mellitus Inadequately Controlled on Metformin Alone")
    kv(doc, "Study Design",    "Randomised, Double-Blind, Double-Dummy, Parallel Group, Active-Controlled, Multicentre")
    kv(doc, "Phase",           "Phase III")
    kv(doc, "Primary Endpoint","Change from baseline in HbA1c at Week 26")
    kv(doc, "Secondary Endpoints","Fasting Plasma Glucose, Body Weight, Systolic BP, Proportion achieving HbA1c < 7.0%, MACE (12-month extended follow-up)")
    kv(doc, "Sample Size",     "540 patients (180 per arm × 3 arms) — 80% power, two-sided α 0.05, expected HbA1c difference 0.7%, SD 1.5%")
    kv(doc, "Duration",        "26 weeks (primary endpoint) + 52-week safety extension")
    kv(doc, "Randomisation",   "1:1:1 — Arm A: EmpaMet 10/500mg, Arm B: EmpaMet 25/1000mg, Arm C: Metformin 1000mg")

    doc.add_paragraph()
    heading(doc, "Part D: Study Sites & Investigators", level=2)
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = 'Table Grid'
    shade_row(t3.rows[0], "D6E4F0")
    for cell, txt in zip(t3.rows[0].cells, ["Site No.", "Institution", "Principal Investigator", "EC Approval"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True
    sites = [
        ("IN-001", "AIIMS New Delhi",                "Dr. Anil Batra, MD, DM Endocrinology",  "EC/AIIMS/2024/019 — 15 Feb 2024 ✓"),
        ("IN-002", "PGI Chandigarh",                 "Dr. Ravi Kant, MD, DM",                  "PENDING — submitted 10 March 2024 ✗"),
        ("IN-003", "JIPMER Puducherry",              "Dr. Subramanian Iyer, MD",               "PENDING — submitted 15 March 2024 ✗"),
        ("IN-004", "KEM Hospital Mumbai",            "Dr. Anand Mehta, MD, DM",                "EC/KEM/2024/34 — 20 Feb 2024 ✓"),
        ("IN-005", "Nizam's Institute Hyderabad",   "Dr. Vijay Viswanathan, MD",              "PENDING — submitted 01 April 2024 ✗"),
        ("IN-006", "Postgraduate Institute Kolkata", "Dr. Sujoy Ghosh, DM",                   "EC/IPGME/2024/112 — 05 March 2024 ✓"),
    ]
    for row in sites:
        table_row(t3, row)

    doc.add_paragraph()
    heading(doc, "Part E: Document Checklist", level=2)
    t4 = doc.add_table(rows=1, cols=3)
    t4.style = 'Table Grid'
    shade_row(t4.rows[0], "D6E4F0")
    for cell, txt in zip(t4.rows[0].cells, ["Document", "Status", "Remarks"]):
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True
    checklist = [
        ("Form CT-04 (Application Form)",              "✓ SUBMITTED",  "Signed by CEO and Head RA"),
        ("Clinical Protocol v3.2",                     "✓ SUBMITTED",  "Dated 15 March 2024, 148 pages"),
        ("Investigator's Brochure v5.0",               "✓ SUBMITTED",  "Dated January 2024, 312 pages"),
        ("Informed Consent Form (English)",             "✓ SUBMITTED",  ""),
        ("Informed Consent Form (Hindi)",               "✓ SUBMITTED",  ""),
        ("ICF in regional languages",                   "✗ NOT SUBMITTED","Tamil, Telugu, Bengali ICFs absent — 3 sites require local language"),
        ("Patient Information Sheet",                   "✓ SUBMITTED",  "English and Hindi versions"),
        ("EC Approval — AIIMS Delhi",                   "✓ SUBMITTED",  "EC/AIIMS/2024/019"),
        ("EC Approval — PGI Chandigarh",                "✗ PENDING",    "Expected April 2024"),
        ("EC Approval — JIPMER",                        "✗ PENDING",    "Expected April 2024"),
        ("EC Approval — KEM Mumbai",                    "✓ SUBMITTED",  "EC/KEM/2024/34"),
        ("EC Approval — Nizam's Institute",             "✗ PENDING",    "Submitted 01 April 2024"),
        ("EC Approval — IPGME Kolkata",                 "✓ SUBMITTED",  "EC/IPGME/2024/112"),
        ("CV of Principal Investigators",               "✓ SUBMITTED",  "All 6 PIs — GCP certified"),
        ("GCP Certificates — All Investigators",        "✓ SUBMITTED",  "Valid certificates for all sites"),
        ("Phase I Data (Bioavailability Study)",        "✓ SUBMITTED",  "Study Sun-EMP-BA-001, n=36"),
        ("Phase II Data (Efficacy/Dose Selection)",     "✓ SUBMITTED",  "Study Sun-EMP-P2-001, n=180, 12 weeks"),
        ("Statistical Analysis Plan (SAP)",             "✗ NOT SUBMITTED","Not included — critical gap"),
        ("Insurance / Indemnity Certificate",           "✓ SUBMITTED",  "New India Assurance, ₹25 lakhs per subject"),
        ("Financial Disclosure Forms — All PIs",        "✓ SUBMITTED",  "No conflicts declared"),
        ("Drug Supply Arrangement",                     "✓ SUBMITTED",  "Qualified Manufacturer: Sun Pharma Silvassa"),
        ("GMP Certificate — Manufacturing Site",        "✓ SUBMITTED",  "Valid to December 2025"),
        ("Bioanalytical Method Validation Report",      "✓ SUBMITTED",  "LC-MS/MS validated method"),
        ("Risk Management Plan",                        "✗ NOT SUBMITTED","Required under NDCT Rules 2019"),
        ("Paediatric Investigation Plan / Waiver",      "✓ SUBMITTED",  "Waiver granted — adult indication only"),
        ("Free Sale Certificate (Germany)",             "✓ SUBMITTED",  "Empagliflozin approved in EU since 2014"),
    ]
    for row in checklist:
        table_row(t4, row)

    doc.add_paragraph()
    heading(doc, "Part F: Previous Regulatory Interactions", level=2)
    para(doc, "Pre-IND meeting with CDSCO held on 14 September 2023. CDSCO indicated that Phase I bridging study in Indian population would be required given known PK differences in Asian vs Caucasian populations. Bridging study completed (Study Sun-EMP-BA-001, n=36 Indian healthy volunteers) — data submitted. CDSCO confirmed adequacy via email dated 15 February 2024.")

    heading(doc, "Part G: Benefit-Risk Assessment", level=2)
    para(doc, "Empagliflozin monotherapy has demonstrated HbA1c reduction of 0.7–0.8% in Indian Phase III data (EMPA-REG India subset). Metformin remains first-line standard of care. The FDC addresses the common clinical problem of pill burden reducing adherence in T2DM patients on multiple medications. EMPA-REG OUTCOME trial demonstrated 38% reduction in cardiovascular death — this benefit-risk is considered strongly positive for the target population (T2DM with inadequate glycaemic control on metformin). Primary risk: Genital mycotic infections (9.7%), UTI (8.1%), DKA (rare, 0.1%).")

    rule(doc)
    para(doc, "Declaration: I/We hereby declare that the information provided in this application and accompanying documents is true, complete and accurate to the best of my/our knowledge.", bold=True)
    para(doc, "Signature: Dr. Anil Varghese, VP Clinical Development | Date: 01 April 2024 | Sun Pharmaceutical Industries Ltd.", italic=True, size=9)

    path = os.path.join(OUT, "04_SUGAM_CT_application_COMPLETENESS_CHECK.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 5a & 5b: Protocol v1 and v2 (Document Comparison)
# ─────────────────────────────────────────────────────────────────────────────

def make_protocol_docs():
    for version, data in [("v2", False), ("v3", True)]:
        doc = new_doc()
        ver_num = "2.0" if not data else "3.0"
        date = "15 January 2024" if not data else "10 March 2024"

        heading(doc, f"CLINICAL STUDY PROTOCOL — AMENDMENT {ver_num}", level=1, color=(0, 61, 122))
        kv(doc, "Protocol No.",   "PSI-CV-2023-01")
        kv(doc, "Version",        ver_num)
        kv(doc, "Date",           date)
        kv(doc, "Drug",           "Cardivex 10mg" if not data else "Cardivex 10mg / 20mg (Dose Escalation Added)")
        kv(doc, "Sponsor",        "Pharma Solutions India Pvt. Ltd., Mumbai")
        rule(doc)

        heading(doc, "1. Study Objectives", level=2)
        heading(doc, "1.1 Primary Objective", level=3)
        if not data:
            para(doc, "To evaluate the efficacy of Cardivex 10mg versus placebo in reducing major adverse cardiovascular events (MACE — defined as cardiovascular death, non-fatal MI, or non-fatal stroke) over 12 months of treatment in adults aged 18–65 years with established cardiovascular disease and LDL-C ≥ 130 mg/dL despite statin therapy.")
        else:
            para(doc, "To evaluate the efficacy of Cardivex 10mg and 20mg versus placebo in reducing major adverse cardiovascular events (MACE — defined as cardiovascular death, non-fatal MI, non-fatal stroke, or hospitalisation for unstable angina) over 18 months of treatment in adults aged 18–70 years with established cardiovascular disease or high cardiovascular risk and LDL-C ≥ 100 mg/dL despite statin therapy.")
            para(doc, "[AMENDMENT NOTE v3.0: MACE definition expanded to include hospitalisation for unstable angina (4-point MACE). Age range extended to 70 years based on regulatory feedback. LDL threshold reduced from 130 to 100 mg/dL to increase eligible population. Duration extended from 12 to 18 months to capture mortality benefit.]", italic=True, color=(139, 0, 0))

        heading(doc, "1.2 Secondary Objectives", level=3)
        sec_obj_v2 = [
            "Change from baseline in LDL-C at 12 months",
            "Change from baseline in total cholesterol, HDL-C, and TG at 12 months",
            "Incidence of individual MACE components",
            "All-cause mortality",
            "Safety and tolerability (AEs, SAEs, laboratory parameters, vital signs)",
        ]
        sec_obj_v3 = [
            "Change from baseline in LDL-C at 18 months",
            "Change from baseline in total cholesterol, HDL-C, TG, Lp(a) at 18 months [Lp(a) added in v3.0]",
            "Incidence of individual 4-point MACE components",
            "All-cause mortality",
            "Hospitalisation for heart failure (new endpoint added v3.0)",
            "Patient-reported outcomes — SF-36 Quality of Life (new endpoint added v3.0)",
            "Safety and tolerability (AEs, SAEs, laboratory parameters, vital signs)",
            "Renal composite endpoint: ≥40% decline in eGFR, renal death, or ESKD (new endpoint added v3.0)",
        ]
        for obj in (sec_obj_v3 if data else sec_obj_v2):
            doc.add_paragraph(obj, style='List Bullet')

        heading(doc, "2. Study Design", level=2)
        kv(doc, "Design",       "Randomised, Double-Blind, Placebo-Controlled, Parallel-Group, Multicentre")
        kv(doc, "Phase",        "Phase III")
        kv(doc, "Treatment Arms","Arm A: Cardivex 10mg OD + Placebo (for 20mg)" if data else "Arm A: Cardivex 10mg OD")
        if data:
            kv(doc, "",         "Arm B: Cardivex 20mg OD + Placebo (for 10mg) [NEW IN v3.0]")
        kv(doc, "",             "Arm C: Placebo OD" if data else "Arm B: Placebo OD")
        kv(doc, "Duration",     "18 months (primary) + 6 months safety follow-up" if data else "12 months (primary) + 3 months safety follow-up")
        kv(doc, "Randomisation","1:1:1 (Cardivex 10mg : Cardivex 20mg : Placebo)" if data else "1:1 (Cardivex 10mg : Placebo)")

        heading(doc, "3. Study Population", level=2)
        heading(doc, "3.1 Inclusion Criteria", level=3)
        inc_v2 = [
            "Age 18–65 years at time of screening",
            "Established cardiovascular disease (prior MI, stroke, or symptomatic PAD)",
            "LDL-C ≥ 130 mg/dL despite statin therapy for ≥ 3 months",
            "On stable background statin therapy for ≥ 3 months",
            "Willing and able to provide informed consent",
        ]
        inc_v3 = [
            "Age 18–70 years at time of screening [AMENDED from 65 in v3.0]",
            "Established cardiovascular disease (prior MI, stroke, or symptomatic PAD) OR high cardiovascular risk (diabetes + ≥ 2 risk factors) [AMENDED in v3.0]",
            "LDL-C ≥ 100 mg/dL despite statin therapy for ≥ 3 months [AMENDED from 130 in v3.0]",
            "On stable background statin therapy for ≥ 3 months",
            "eGFR ≥ 30 mL/min/1.73m² (renal criterion added in v3.0 for renal endpoint validity)",
            "Willing and able to provide informed consent",
        ]
        for c in (inc_v3 if data else inc_v2):
            doc.add_paragraph(c, style='List Bullet')

        heading(doc, "3.2 Exclusion Criteria", level=3)
        excl_v2 = [
            "Severe renal impairment (eGFR < 30 mL/min/1.73m²)",
            "Active liver disease or ALT/AST > 3× ULN",
            "Known statin intolerance or hypersensitivity to any component",
            "NYHA Class IV heart failure",
            "Life expectancy < 12 months",
            "Current participation in another interventional trial",
            "Pregnancy or breastfeeding",
        ]
        excl_v3 = [
            "Severe renal impairment (eGFR < 15 mL/min/1.73m²) [AMENDED from <30 in v3.0 — to allow CKD Stage 3-4 patients for renal endpoint]",
            "Active liver disease or ALT/AST > 3× ULN",
            "Known statin intolerance or previous rhabdomyolysis on any statin [AMENDED — rhabdomyolysis exclusion added post SAE-2024-MH-0341]",
            "NYHA Class IV heart failure",
            "Life expectancy < 18 months [AMENDED from 12 months]",
            "Current participation in another interventional trial",
            "Pregnancy or breastfeeding",
            "Organ transplant recipients on immunosuppressants [NEW EXCLUSION — drug interaction risk v3.0]",
            "Active malignancy within 5 years [NEW EXCLUSION v3.0]",
        ]
        for c in (excl_v3 if data else excl_v2):
            doc.add_paragraph(c, style='List Bullet')

        heading(doc, "4. Sample Size", level=2)
        if not data:
            para(doc, "Sample Size Calculation: 300 patients total (150 per arm). Based on expected 15% reduction in MACE at 12 months (treatment arm 12% vs placebo 14.5%), 80% power, two-sided α 0.05, 20% dropout rate. Event-driven analysis: 60 primary events required.")
        else:
            para(doc, "Sample Size Calculation (AMENDED v3.0): 450 patients total (150 per arm × 3 arms). Based on expected 18% reduction in 4-point MACE at 18 months (treatment arm 10% vs placebo 12.5%), 85% power (increased from 80%), two-sided α 0.05, 25% dropout rate (increased from 20% due to extended duration). Event-driven analysis: 90 primary events required.\n[RATIONALE FOR INCREASE: Additional arm, extended duration, expanded MACE definition, and higher power requirement per DSMB recommendation at interim analysis dated 20 February 2024.]")

        heading(doc, "5. Dosage & Administration", level=2)
        if not data:
            para(doc, "Cardivex 10mg: One tablet orally once daily, in the morning with food. Placebo: One matching tablet orally once daily. Treatment Duration: 12 months. Dose modifications: Not permitted. Study drug discontinuation required for: CK > 10× ULN, ALT/AST > 5× ULN, or any SAE assessed as related to study drug.")
        else:
            para(doc, "Cardivex 10mg Arm: One 10mg tablet + one matching placebo (for 20mg) orally once daily, in the morning with food.\nCardivex 20mg Arm (NEW v3.0): One 20mg tablet + one matching placebo (for 10mg) orally once daily, in the morning with food. Dose escalation from 10mg to 20mg at Week 4 visit if 10mg tolerated (CK < 5× ULN, no myalgia).\nPlacebo Arm: Two matching placebo tablets once daily.\nTreatment Duration: 18 months. \nSafety stopping rules: Permanent discontinuation required if CK > 10× ULN with symptoms OR CK > 40× ULN (asymptomatic). [UPDATED per rhabdomyolysis SAE — earlier stopping threshold added.]")

        heading(doc, "6. Statistical Analysis Plan", level=2)
        if not data:
            para(doc, "Primary analysis: ANCOVA with baseline HbA1c and randomisation strata as covariates. ITT population. Missing data: Multiple imputation (MI) using MICE algorithm. Multiplicity: Hierarchical testing — primary endpoint first, then secondary endpoints in pre-specified order. Interim analysis: One planned at 50% events by independent DSMB.")
        else:
            para(doc, "[AMENDED v3.0] Primary analysis: Cox proportional hazards model with treatment arm, country, baseline LDL-C stratum, and cardiovascular risk category as covariates. Time-to-first-event analysis. ITT and Per Protocol populations both reported. Missing data: Multiple imputation with pattern mixture model sensitivity analysis. Multiplicity: Hierarchical gatekeeping procedure — 3 arms require Bonferroni-Holm correction for pairwise comparisons. Interim analyses: Two planned — at 30% and 60% events by independent DSMB with O'Brien-Fleming spending function. [Statistical Analysis Plan v3.0 attached as Appendix G — previously absent in v2.0 submission, now included as separate document per CDSCO request dated 12 February 2024]")

        path = os.path.join(OUT, f"05{'b' if data else 'a'}_protocol_{version}_COMPARE_THESE.docx")
        doc.save(path)
        print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 6: GMP Inspection Notes (Inspection Report Generator)
# ─────────────────────────────────────────────────────────────────────────────

def make_inspection_notes_doc():
    doc = new_doc()

    heading(doc, "GMP INSPECTION — FIELD NOTES", level=1, color=(100, 0, 0))
    para(doc, "CONFIDENTIAL — For Official Use Only | Inspector Working Papers", bold=True, color=(180, 0, 0))
    rule(doc)

    kv(doc, "Inspection Type",    "Announced GMP Inspection — Manufacturing Site")
    kv(doc, "Facility",           "Shivalik Pharmaceuticals Pvt. Ltd., Unit II")
    kv(doc, "Address",            "Plot 24-B, Phase II, GIDC Industrial Estate, Ankleshwar, Gujarat — 393002")
    kv(doc, "Drug Licence No.",   "GUJ/MFG/2009/1204")
    kv(doc, "Inspection Dates",   "10–12 April 2024")
    kv(doc, "Lead Inspector",     "Mr. Dinesh Sharma, Deputy Drugs Controller (I), Zone III, CDSCO")
    kv(doc, "Co-Inspector",       "Dr. Kavita Rao, Assistant Drugs Controller, Gujarat FDA")
    kv(doc, "Company Representative","Mr. Rajiv Mehta, Plant Head | Mr. Suresh Patel, QA Head")
    doc.add_paragraph()

    heading(doc, "Day 1 — 10 April 2024 Observations", level=2)

    heading(doc, "Area 1: Warehouse / Storage (Raw Material and Finished Goods)", level=3)
    storage_obs = [
        "Cold storage unit (2–8°C) — ambient temperature reading was 11°C at 10:30 AM on day of inspection. Temperature logger downloaded — shows excursion to 13.4°C for 6 hours on 05 April 2024. No deviation report raised, no corrective action taken. Batch RM-API-0234 (Cefuroxime Axetil, 200 kg) stored during this period — no quarantine or retest initiated.",
        "General warehouse temperature — target 15–25°C. Measured 29.8°C at 11:00 AM. HVAC maintenance record requested — last service August 2022, filter change January 2023. No validation of HVAC performance in past 18 months.",
        "Humidity in finished goods store: 72% RH (limit: NMT 60% RH). Desiccant bags not changed — saturated condition noted.",
        "Rejected material area: No clear physical segregation from quarantine materials. Rejected batch RFG-2024-0089 (Metformin tablets, 50,000 units, rejected for failing dissolution) stored alongside quarantined incoming RM. Labels not adequately visible — double-sided tape failing.",
        "Narcotics cabinet — lock intact but key was lying on QC Manager's table (Room QC-Lab-01). Access record for previous 30 days not available.",
        "GDP records incomplete — dispatch records for January 2024 missing Chain of Custody documentation for 3 consignments to Maharashtra.",
    ]
    for i, obs in enumerate(storage_obs, 1):
        para(doc, f"OBS-S{i:02d}: {obs}")

    heading(doc, "Area 2: Production Floor — Solid Dosage (Tablet Manufacturing)", level=3)
    prod_obs = [
        "Personnel hygiene: 3 production staff observed in granulation area without gloves. 1 operator observed touching tablet compression output with bare hands. Hair nets worn incorrectly — ears exposed for 2 operators. Senior supervisor (Mr. Govind Sharma) present in area, did not intervene.",
        "Cleaning validation: Shared equipment — V-blender used for both penicillin (Ampicillin) and non-penicillin products. Last cleaning validation study conducted 2019. Visual inspection only between campaigns. Swab test not performed before changeover on 09 April 2024 (previous product: Ampicillin 250mg, current: Metformin 500mg). MAJOR cross-contamination risk for penicillin-sensitive patients.",
        "Batch Manufacturing Record for Batch MF-TAB-2024-089 (Metformin 500mg, 1,00,000 tablets): Page 4 of 12 (In-process checks section) — QA supervisor signature absent. Granulation yield recorded as 98.4% but raw material issue records show only sufficient for 96.2% yield — discrepancy unresolved. Line clearance certificate signed but date field blank.",
        "Tablet compression machine TC-04 — oil stains on hopper and die table. Last cleaned according to log: 07 April 2024. Cleaning record signed by operator but verified signature of QC person missing.",
        "Environmental monitoring in production area: Settle plates not available for Grade C area. Last environmental monitoring: January 2024. Alert and action limits not defined in SOP.",
        "In-process testing lab: UV-Vis spectrophotometer (Model Shimadzu UV-1800, Asset No. QC-INST-012) — last calibrated October 2022. Calibration SOP requires annual calibration. Calibration sticker expired. Results from this instrument used for in-process dissolution testing of Batch MF-TAB-2024-089.",
    ]
    for i, obs in enumerate(prod_obs, 1):
        para(doc, f"OBS-P{i:02d}: {obs}")

    heading(doc, "Day 2 — 11 April 2024 Observations", level=2)

    heading(doc, "Area 3: Quality Control Laboratory", level=3)
    qc_obs = [
        "HPLC system (Agilent 1260, Asset QC-HPLC-002): OQ/PQ last performed March 2022. Computer access controls — shared login 'qcanalyst1' used by at least 4 analysts simultaneously. Audit trail review shows backdated entries for three FG release tests dated 15–17 March 2024, entered on 19 March 2024. Test entries not matching original lab notebooks.",
        "Reference standard management: Working standard of Cefuroxime Axetil (RS-CEF-024) — characterisation certificate expired 31 December 2023. Still in active use as per lab register entry dated 08 April 2024. No retest performed.",
        "Out-of-Specification (OOS) investigation log: 3 OOS results recorded (Ref: OOS-2024-003, OOS-2024-007, OOS-2024-009). OOS-2024-003 investigation closed in 2 days — atypical (normal is 10–15 days). Investigation report reviewed: Phase I (lab error investigation) — analyst attributed to 'pipetting error'. No Phase II investigation (full-scale investigation) conducted. Batch released to market.",
        "Water system: Purified Water generation system (WFI plant) — last full validation 2019. Chemical monitoring done weekly, but microbial testing frequency reduced from daily to twice weekly since February 2024 due to 'manpower shortage'. WFI not in scope for this inspection but PW is. PW loop microbial result for 25 March 2024: 98 CFU/mL (Alert limit: 50 CFU/mL, Action limit: 100 CFU/mL). No deviation raised.",
        "Stability chambers: 40°C/75%RH chamber (long-term accelerated stability) — temperature alarm triggered at 43.2°C on 28 February 2024 for 14 hours. No investigation. Samples of 3 products in long-term stability study may be compromised.",
    ]
    for i, obs in enumerate(qc_obs, 1):
        para(doc, f"OBS-Q{i:02d}: {obs}")

    heading(doc, "Area 4: Documentation & Quality Systems", level=3)
    doc_obs = [
        "Standard Operating Procedures — SOP for 'Handling of Rejected Materials' (SOP-QA-012): last review 2020. Current revision due date: January 2023 — overdue by 15 months. No periodic review records.",
        "Change Control system: Review of change control log shows 12 changes implemented in 2023 — 4 with risk categorisation 'Major' (including equipment replacement and starting material supplier change). CDSCO notification not done for any of these changes as required under Schedule M Rule 71(2).",
        "Vendor qualification: New API supplier (Shree Ram Chemicals) added in November 2023 for Metformin. Approved vendor list updated. Audit of vendor — not conducted. Certificate of Analysis accepted without independent testing for first 3 batches.",
        "Product Quality Review (PQR) for FY 2022-23: Not conducted. PQR for FY 2023-24: In progress, projected completion June 2024 (already 3 months overdue for FY23 under Schedule M requirement).",
        "Recall procedure: SOP available. Mock recall exercise — last conducted August 2021. CDSCO requires annual mock recall. Missed 2022 and 2023.",
    ]
    for i, obs in enumerate(doc_obs, 1):
        para(doc, f"OBS-D{i:02d}: {obs}")

    heading(doc, "Day 3 — 12 April 2024 Observations (Closing)", level=2)

    heading(doc, "Area 5: Pest Control and Facility Infrastructure", level=3)
    facility_obs = [
        "Pest control: Last treatment 14 November 2023 — 5 months prior. Annual contract in place with Pepilite Pest Control. However, contract requires quarterly treatment. Two intermediate treatments (February and April 2024) not completed. Rodent bait station check records for March 2024 — not available.",
        "Drainage in packaging area: Open drain observed adjacent to blister packaging line. Risk of contamination from water splash. Cover plate missing since 'at least December 2023' as per production supervisor.",
        "Air handling unit — packaging hall: HEPA filter (H14 grade) last replaced January 2022. Differential pressure across filter: 280 Pa (change limit: 250 Pa — already exceeded). Filter overdue for replacement. Non-viable particle counts not available for this area since September 2023.",
        "Manufacturing area entry airlocks: Interlock on airlock door to packaging hall (PAC-AIR-02) non-functional. Both doors can be opened simultaneously — defeating airlock purpose. Reported to maintenance on 22 March 2024 per maintenance log but not yet repaired.",
    ]
    for i, obs in enumerate(facility_obs, 1):
        para(doc, f"OBS-F{i:02d}: {obs}")

    heading(doc, "Summary of Critical Concerns Noted by Inspectors", level=2)
    para(doc, "The following observations are considered CRITICAL and require immediate CAPA:", bold=True, color=(180, 0, 0))
    critical = [
        "Cross-contamination risk — penicillin/non-penicillin shared equipment without validated cleaning (OBS-P02)",
        "Backdated audit trail entries in HPLC system — potential data integrity violation (OBS-Q01)",
        "OOS investigation not completed per regulatory requirement — batch released to market (OBS-Q03)",
        "Expired reference standard used for FG release testing (OBS-Q02)",
        "Cold chain excursion not investigated — API batch of unknown quality in supply chain (OBS-S01)",
    ]
    for c in critical:
        doc.add_paragraph(c, style='List Bullet')

    rule(doc)
    para(doc, "These notes are preliminary and subject to final report preparation. Company response expected within 15 working days of formal report issue.", italic=True, size=9)
    para(doc, f"Inspector: Mr. Dinesh Sharma | Signature: [Signed] | Date: 12 April 2024", size=9)
    para(doc, f"Co-Inspector: Dr. Kavita Rao | Signature: [Signed] | Date: 12 April 2024", size=9)

    path = os.path.join(OUT, "06_GMP_inspection_notes_GENERATE_REPORT.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# Run all
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating test documents...\n")
    make_anonymisation_doc()
    make_sae_narration_doc()
    make_meeting_transcript_doc()
    make_sugam_application_doc()
    make_protocol_docs()
    make_inspection_notes_doc()
    print(f"\nAll documents saved to: {OUT}")
    print("\nUsage guide:")
    print("  01_patient_case_record*.docx      → Anonymisation page (try both modes)")
    print("  02_SAE_narration*.docx            → Summarisation (select: SAE Case Narration)")
    print("  03_CDSCO_meeting_transcript*.docx → Summarisation (select: Meeting Transcript)")
    print("  04_SUGAM_CT_application*.docx     → Completeness (Clinical Trial Application checklist)")
    print("  04_SUGAM_CT_application*.docx     → Summarisation (select: SUGAM Application)")
    print("  05a_protocol_v2*.docx + 05b*      → Completeness → Document Comparison tab")
    print("  06_GMP_inspection_notes*.docx     → Inspection Report (select: GMP Inspection)")
